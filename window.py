import ast
import csv
import itertools
import math
import os
import re
import sys
from collections import Counter
from collections import defaultdict
from typing import Dict
import pickle
import matplotlib.pyplot as plt

import common
from UI import Ui_APTMainWindow, Ui_InputElementTable, Ui_PeriodicTable, Ui_MonoLayer, Ui_DecomposeList, \
    Ui_AbstractLayer, Ui_SRO, Ui_CompositionMap

import docx
import numpy as np
import pandas as pd
from IPython.external.qt_for_kernel import QtCore
from PyQt5 import QtGui
from PyQt5.QtCore import QEvent, Qt
from PyQt5.QtGui import QCursor, QStandardItemModel, QStandardItem
from PyQt5.QtWidgets import QMainWindow, QFileDialog, QDialog, QVBoxLayout, QTableWidgetItem, \
    QHeaderView, QLineEdit, QLabel, QDialogButtonBox
from ase.io import read
from ase.neighborlist import NeighborList, NewPrimitiveNeighborList
from matplotlib.backends.backend_qt5 import NavigationToolbar2QT
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.figure import Figure
from natsort import index_natsorted, order_by_index
from scipy import special
from scipy.signal import savgol_filter
from scipy.spatial import ConvexHull, Delaunay
from silx.gui.widgets.PeriodicTable import PeriodicTable
from sklearn.cluster import DBSCAN
from sklearn.neighbors import KDTree

pd.options.mode.chained_assignment = None
pd.set_option('display.width', 320)
np.set_printoptions(linewidth=320)
pd.set_option('display.max_columns', 10)
pickle.HIGHEST_PROTOCOL = 4

# Functions defined in the common class
show_message = common.show_message

# maximum number of ions can be specified here
max_ions = 50
# The dictionary of ions for global use and update. Each ion will be one element_dict.
element_dict = {'ion': [], 'num': [], 'mass': [], 'charge': []}
cutoff_dict = {'peak_MNRatio': float, 'peak_width': float, 'cutoff_bin': float, 'cutoff_height': int,
               'cutoff_width': int}
# list of element_dict corresponding to row number
element_dict_array: Dict[int, dict] = {}
# list of cutoff values corresponding to row number
cutoff_dict_array: Dict[int, dict] = {}


# The following class is used to visualize the dataframe inside the main window. The column size could be readjusted.
# reference: https://stackoverflow.com/questions/44603119/how-to-display-a-pandas-data-frame-with-pyqt5-pyside2
# Does not inherit any UI files
class DataFrameModel(QtCore.QAbstractTableModel):
    DtypeRole = QtCore.Qt.UserRole + 1000
    ValueRole = QtCore.Qt.UserRole + 1001

    def __init__(self, df=pd.DataFrame(), parent=None):
        super(DataFrameModel, self).__init__(parent)
        self._dataframe = df

    def setDataFrame(self, dataframe):
        self.beginResetModel()
        self._dataframe = dataframe.copy()
        self.endResetModel()

    def dataFrame(self):
        return self._dataframe

    dataFrame = QtCore.pyqtProperty(pd.DataFrame, fget=dataFrame, fset=setDataFrame)

    @QtCore.pyqtSlot(int, QtCore.Qt.Orientation, result=str)
    def headerData(self, section: int, orientation: QtCore.Qt.Orientation, role: int = QtCore.Qt.DisplayRole):
        if role == QtCore.Qt.DisplayRole:
            if orientation == QtCore.Qt.Horizontal:
                return self._dataframe.columns[section]
            else:
                return str(self._dataframe.index[section])
        return QtCore.QVariant()

    def rowCount(self, parent=QtCore.QModelIndex()):
        if parent.isValid():
            return 0
        return len(self._dataframe.index)

    def columnCount(self, parent=QtCore.QModelIndex()):
        if parent.isValid():
            return 0
        return self._dataframe.columns.size

    def data(self, index, role=QtCore.Qt.DisplayRole):
        if not index.isValid() or not (0 <= index.row() < self.rowCount() and 0 <= index.column() < self.columnCount()):
            return QtCore.QVariant()
        row = self._dataframe.index[index.row()]
        col = self._dataframe.columns[index.column()]
        dt = self._dataframe[col].dtype

        val = self._dataframe.iloc[row][col]
        if role == QtCore.Qt.DisplayRole:
            return str(val)
        elif role == DataFrameModel.ValueRole:
            return val
        if role == DataFrameModel.DtypeRole:
            return dt
        return QtCore.QVariant()

    def roleNames(self):
        roles = {
            QtCore.Qt.DisplayRole: b'display',
            DataFrameModel.DtypeRole: b'dtype',
            DataFrameModel.ValueRole: b'value'
        }
        return roles

    def sort(self, column, order):
        self.layoutAboutToBeChanged.emit()
        if order == 0:
            self._dataframe = self._dataframe.reindex(index=order_by_index(self._dataframe.index, index_natsorted(
                eval('self._dataframe.%s' % (list(self._dataframe.columns)[column])))))
        else:
            self._dataframe = self._dataframe.reindex(index=order_by_index(self._dataframe.index, reversed(
                index_natsorted(eval('self._dataframe.%s' % (list(self._dataframe.columns)[column]))))))

        self._dataframe.reset_index(inplace=True, drop=True)
        self.setDataFrame(self._dataframe)
        self.layoutChanged.emit()


# Class used to plot and display histogram
# Does not inherit any UI files
class HistogramWindow(QDialog):
    def __init__(self, subset_df_apt_hist, parent=None):
        super(HistogramWindow, self).__init__(parent)

        # a figure instance to plot on
        self.figure = Figure()
        self.subset_df_apt_hist = subset_df_apt_hist

        # this is the Canvas Widget that displays the `figure`
        # it takes the `figure` instance as a parameter to __init__
        self.canvas = FigureCanvasQTAgg(self.figure)

        # this is the Navigation widget
        # it takes the Canvas widget and a parent
        self.toolbar = NavigationToolbar2QT(self.canvas, self)

        # set the layout
        layout = QVBoxLayout()
        layout.addWidget(self.toolbar)
        layout.addWidget(self.canvas)
        self.setLayout(layout)
        self.plot()

    def plot(self):
        # create an axis
        ax = self.figure.add_subplot(111)

        # discards the old graph
        ax.clear()

        # plot data
        ax.title.set_text('APT Spectrum')
        ax.set_xlabel('MN_Ratio')
        ax.set_ylabel('Counts')
        ax.bar(self.subset_df_apt_hist.bin_lower, self.subset_df_apt_hist.freq,
               width=self.subset_df_apt_hist.bin_upper - self.subset_df_apt_hist.bin_lower, ec="k", align="edge")

        # refresh canvas
        self.canvas.draw()


# The class shows an on-the-top dialog that can input num of elements and mass (different from default for isotope).
# Does not inherit any UI files
class NumberAndMass(QDialog):
    def __init__(self, parent=None):
        super(NumberAndMass, self).__init__(parent)

        mainLayout = QVBoxLayout()
        self.lineedit, self.lineedit2 = QLineEdit(), QLineEdit()
        self.label, self.label2 = QLabel(), QLabel()
        self.btns = QDialogButtonBox()
        self.btns.setStandardButtons(QDialogButtonBox.Cancel | QDialogButtonBox.Ok)
        self.label.setText("Enter No: of atoms (default 1)")
        self.label2.setText("Enter Atomic Mass (default as filled)")
        mainLayout.addWidget(self.label)
        mainLayout.addWidget(self.lineedit)
        mainLayout.addWidget(self.label2)
        mainLayout.addWidget(self.lineedit2)
        mainLayout.addWidget(self.btns)
        self.btns.accepted.connect(self.accept)
        self.btns.rejected.connect(self.reject)
        self.setLayout(mainLayout)


# The class inherits a custom UI layout and has facility to view periodic table, add elements into the table
# Inherits from Ui_PeriodicTable
class PeriodicTableCustom(Ui_PeriodicTable.Ui_Form, QDialog):
    def __init__(self, parent=None):
        super(PeriodicTableCustom, self).__init__(parent)
        self.setupUi(self)

        # inputDialog is used to get no: of atoms
        self.inputDialog = NumberAndMass()
        self.inputDialog.setWindowTitle('Input')
        self.lineEdit.setText("0")
        self.curr_row = None

        self.ptable = PeriodicTable(self.widget, selectable=False)
        self.ptable.sigElementClicked.connect(self.click_table)
        self.inputDialog.setWindowModality(QtCore.Qt.ApplicationModal)
        self.pushButton_2.clicked.connect(self.delete)
        self.pushButton_5.clicked.connect(self.emit_charge)
        self.pushButton_5.clicked.connect(self.accept)
        self.pushButton_4.clicked.connect(self.reject)

    def delete(self):
        text = ""
        if len(element_dict['ion']) > 0:
            del element_dict['ion'][-1]
            del element_dict['num'][-1]
            for i in range(len(element_dict['ion'])):
                text = text + str(element_dict['ion'][i]) + common.subscript(str(element_dict['num'][i]))
        self.lineEdit_3.setText(text)

    def emit_charge(self):
        text = self.lineEdit.text()
        if '-' in text:
            sign = '-'
        else:
            sign = '+'
        text_num = re.sub('\D', '', text)
        text = text_num + sign

        return text

    def click_table(self, item):
        self.showdialogTOP()
        self.inputDialog.lineedit2.setText(str(item.mass))
        self.inputDialog.lineedit.setText(str(1))
        ok = self.inputDialog.exec_()

        elem_num = self.inputDialog.lineedit.text()
        elem_mass = self.inputDialog.lineedit2.text()

        if ok:
            if len(element_dict['ion']) == 0:
                self.lineEdit_3.setText("")

            element_dict['ion'].append(str(item.symbol))
            element_dict['num'].append(elem_num)
            element_dict['mass'].append(elem_mass)
            element_dict['charge'] = [self.emit_charge()]

            text = ""
            for i in range(len(element_dict['ion'])):
                text = text + str(element_dict['ion'][i]) + common.subscript(str(element_dict['num'][i]))

            self.lineEdit_3.setText(text)
            element_dict_array[str(self.curr_row)] = dict()
            element_dict_array[str(self.curr_row)]['ion'] = element_dict['ion']
            element_dict_array[str(self.curr_row)]['num'] = element_dict['num']
            element_dict_array[str(self.curr_row)]['mass'] = element_dict['mass']
            element_dict_array[str(self.curr_row)]['charge'] = element_dict['charge']

    # The following makes sure the input dialogue to enter no: of atoms stays on top and on cursor location
    def showdialogTOP(self):
        self.inputDialog.move(QCursor.pos())
        self.inputDialog.show()
        self.inputDialog.raise_()

    def showEvent(self, event):
        if len(element_dict['ion']) == 0:
            self.lineEdit_3.setText("")
        text = ""
        for i in range(len(element_dict['ion'])):
            text = text + str(element_dict['ion'][i]) + common.subscript(str(element_dict['num'][i]))

        self.lineEdit_3.setText(text)
        event.accept()


# The class inherits a custom UI layout and has facility to complete the input elements table as well as cutoff values
# Inherits from Ui_InputElementTable
class InputElementTable(Ui_InputElementTable.Ui_Form, QDialog):
    def __init__(self, parent=None):
        super(InputElementTable, self).__init__(parent)
        self.setupUi(self)

        self.tableWidget.setRowCount(max_ions)
        self.tableWidget.setColumnCount(6)
        self.tableWidget.setHorizontalHeaderItem(0, QTableWidgetItem("ION"))
        self.tableWidget.setHorizontalHeaderItem(1, QTableWidgetItem("peak_MNRatio(Da)"))
        self.tableWidget.setHorizontalHeaderItem(2, QTableWidgetItem("MNRatio_width(Da)"))
        self.tableWidget.setHorizontalHeaderItem(3, QTableWidgetItem("cutoff_bin"))
        self.tableWidget.setHorizontalHeaderItem(4, QTableWidgetItem("cutoff_height"))
        self.tableWidget.setHorizontalHeaderItem(5, QTableWidgetItem("cutoff_width"))
        self.PeriodicTableCustom = None
        self.row = 0
        self.col = 0
        self.item = None
        self.cell_clicked = False
        self.buttonBox.accepted.connect(self.submit_close)
        self.buttonBox.rejected.connect(self.reject)

        self.setMinimumSize(500, 500)
        self.setGeometry(QtCore.QRect(417, 220, 666, 653))

        header = self.tableWidget.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.Stretch)
        header.setSectionResizeMode(1, QHeaderView.ResizeToContents)
        header.setSectionResizeMode(2, QHeaderView.ResizeToContents)
        header.setSectionResizeMode(3, QHeaderView.ResizeToContents)
        header.setSectionResizeMode(4, QHeaderView.ResizeToContents)

        self.tableWidget.cellClicked.connect(self.cell_was_clicked)
        self.tableWidget.cellDoubleClicked.connect(self.cell_was_clicked)
        self.pushButton_2.clicked.connect(self.export_table)
        self.pushButton.clicked.connect(self.import_table)
        viewport = self.tableWidget.viewport()
        viewport.installEventFilter(self)

    def saveFileDialog(self):
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        fileName, _ = QFileDialog.getSaveFileName(self, "QFileDialog.getSaveFileName()", "",
                                                  "All Files (*);;csv Files (*.csv)", options=options)
        if fileName:
            return fileName + '.csv'

    def submit(self):
        key = None
        for r in range(max_ions):
            for c in range(1, 4):
                if c == 1:
                    key = 'ion'
                if c == 2:
                    key = 'num'
                if c == 3:
                    key = 'mass'
                if c == 4:
                    key = 'charge'

            for c in range(1, 6):
                if c == 1:
                    key = 'peak_MNRatio'
                if c == 2:
                    key = 'peak_width'
                if c == 3:
                    key = 'cutoff_bin'
                if c == 4:
                    key = 'cutoff_height'
                if c == 5:
                    key = 'cutoff_width'

                if self.tableWidget.item(r, c):
                    cutoff_dict[key] = self.tableWidget.item(r, c).text()
                else:
                    cutoff_dict[key] = None

            cutoff_dict_array[str(r)] = dict()
            cutoff_dict_array[str(r)]['peak_MNRatio'] = cutoff_dict['peak_MNRatio']
            cutoff_dict_array[str(r)]['peak_width'] = cutoff_dict['peak_width']
            cutoff_dict_array[str(r)]['cutoff_bin'] = cutoff_dict['cutoff_bin']
            cutoff_dict_array[str(r)]['cutoff_height'] = cutoff_dict['cutoff_height']
            cutoff_dict_array[str(r)]['cutoff_width'] = cutoff_dict['cutoff_width']

    def export_table(self):
        csv_file, extension = QFileDialog.getSaveFileName(
            self, 'Save File', '.', filter=self.tr("csv file (*.csv)"))
        csv_columns = ['ion', 'num', 'mass', 'charge', 'peak_MNRatio', 'peak_width',
                       'cutoff_bin', 'cutoff_height', 'cutoff_width']

        self.submit()
        list_of_dict = []
        list1 = list(element_dict_array.keys())
        list1 = [int(i) for i in list1]
        list2 = list(cutoff_dict_array.keys())
        list2 = [int(i) for i in list2]
        if min(list1) <= min(list2):
            for key1 in element_dict_array:
                x = element_dict_array[key1]
                if int(key1) in list2:
                    y = cutoff_dict_array[key1]
                else:
                    y = {}
                z = {**x, **y}
                list_of_dict.append(z)

            for key2 in cutoff_dict_array:
                if int(key2) not in list1:
                    x = {}
                    y = cutoff_dict_array[key2]
                    z = {**x, **y}
                    list_of_dict.append(z)
        else:
            for key2 in cutoff_dict_array:
                y = cutoff_dict_array[key2]
                if int(key2) in list1:
                    x = element_dict_array[key2]
                else:
                    x = {}
                z = {**x, **y}
                list_of_dict.append(z)

            for key1 in element_dict_array:
                if int(key1) not in list2:
                    x = element_dict_array[key1]
                    y = {}
                    z = {**x, **y}
                    list_of_dict.append(z)

        try:
            with open(csv_file, 'w', newline="") as csvfile:
                writer = csv.DictWriter(csvfile, fieldnames=csv_columns)
                writer.writeheader()

                for row in list_of_dict:
                    writer.writerow(row)

        except IOError:
            common.show_message("I/O error")

    def import_table(self):
        csv_file, extension = QFileDialog.getOpenFileName(
            self, 'Save File', '.', filter=self.tr("csv file (*.csv)"))

        for value in element_dict.values():
            del value[:]
        for value in cutoff_dict.values():
            del value

        element_dict_array.clear()
        cutoff_dict_array.clear()

        if csv_file:
            try:
                with open(csv_file, "r") as fileInput:
                    for row_num, row in enumerate(csv.reader(fileInput)):
                        if row_num > 0:
                            r = row_num - 1
                            if row[0] != '':
                                element_dict_array[str(r)] = dict()
                                element_dict_array[str(r)]['ion'] = ast.literal_eval(row[0])
                                element_dict_array[str(r)]['num'] = ast.literal_eval(row[1])
                                element_dict_array[str(r)]['mass'] = ast.literal_eval(row[2])
                                element_dict_array[str(r)]['charge'] = ast.literal_eval(row[3])

                            cutoff_dict_array[str(r)] = dict()
                            cutoff_dict_array[str(r)]['peak_MNRatio'] = (row[4])
                            cutoff_dict_array[str(r)]['peak_width'] = (row[5])
                            cutoff_dict_array[str(r)]['cutoff_bin'] = (row[6])
                            cutoff_dict_array[str(r)]['cutoff_height'] = (row[7])
                            cutoff_dict_array[str(r)]['cutoff_width'] = (row[8])

            except IOError:
                common.show_message("I/O error")

        self.refresh_table()

    def refresh_table(self):
        if len(element_dict_array) != 0:
            for key in element_dict_array:
                if eval(key) is not None:
                    row = int(key)
                    value = ''
                    for ii in range(len(element_dict_array[key]['ion'])):
                        value = value + element_dict_array[key]['ion'][ii] + common.subscript(
                            element_dict_array[key]['num'][ii])
                        if ii == len(element_dict_array[key]['ion']) - 1:
                            value = value + '(' + element_dict_array[key]['charge'][0] + ')'

                    item_ion = QTableWidgetItem()
                    item_ion.setText(value)
                    self.tableWidget.setItem(row, 0, item_ion)

        if len(cutoff_dict_array) != 0:
            for key in cutoff_dict_array:
                row = int(key)
                peak_MNRatio = cutoff_dict_array[key]['peak_MNRatio']
                peak_width = cutoff_dict_array[key]['peak_width']
                cutoff_bin = cutoff_dict_array[key]['cutoff_bin']
                cutoff_height = cutoff_dict_array[key]['cutoff_height']
                cutoff_width = cutoff_dict_array[key]['cutoff_width']
                if peak_MNRatio:
                    item_peak = QTableWidgetItem()
                    item_peak.setText(str(peak_MNRatio))
                    self.tableWidget.setItem(row, 1, item_peak)
                if peak_width:
                    item_peak = QTableWidgetItem()
                    item_peak.setText(str(peak_width))
                    self.tableWidget.setItem(row, 2, item_peak)
                if cutoff_bin:
                    item_cutoff_bin = QTableWidgetItem()
                    item_cutoff_bin.setText(str(cutoff_bin))
                    self.tableWidget.setItem(row, 3, item_cutoff_bin)
                if cutoff_height:
                    item_cutoff_height = QTableWidgetItem()
                    item_cutoff_height.setText(str(cutoff_height))
                    self.tableWidget.setItem(row, 4, item_cutoff_height)
                if cutoff_width:
                    item_cutoff_width = QTableWidgetItem()
                    item_cutoff_width.setText(str(cutoff_width))
                    self.tableWidget.setItem(row, 5, item_cutoff_width)

    def showEvent(self, event):
        super(InputElementTable, self).showEvent(event)
        self.refresh_table()

    def submit_close(self):
        self.submit()
        self.accept()
        self.close()

    def cell_was_clicked(self, row, column):
        self.cell_clicked = True
        self.row = row
        self.col = column
        self.item = self.tableWidget.itemAt(row, column)

    def eventFilter(self, source: QtCore.QObject, event: QtCore.QEvent) -> bool:
        if event.type() == QEvent.MouseButtonDblClick:
            for key in element_dict:
                element_dict[key] = []

            headertext = self.tableWidget.horizontalHeaderItem(self.col).text()
            if headertext == 'ION' and self.cell_clicked:
                self.PeriodicTableCustom = PeriodicTableCustom()
                if self.tableWidget.item(self.row, 0):
                    already_ion = self.tableWidget.item(self.row, 0).text()
                    if already_ion:
                        self.PeriodicTableCustom.lineEdit_3.setText(already_ion)

                self.PeriodicTableCustom.curr_row = self.row
                if self.PeriodicTableCustom.exec() == 1:
                    element_dict['charge'] = [self.PeriodicTableCustom.emit_charge()]

                    element_dict_array[str(self.row)] = dict()
                    element_dict_array[str(self.row)]['ion'] = element_dict['ion']
                    element_dict_array[str(self.row)]['num'] = element_dict['num']
                    element_dict_array[str(self.row)]['mass'] = element_dict['mass']
                    element_dict_array[str(self.row)]['charge'] = element_dict['charge']

                    text = ''
                    for i in range(len(element_dict['ion'])):
                        text = text + str(element_dict['ion'][i]) + common.subscript(str(element_dict['num'][i]))
                    if text:
                        charge = str(element_dict['charge'][0])
                        charge = '(' + charge + ')'
                        text = text + charge
                    item = QTableWidgetItem()
                    item.setText(text)
                    self.tableWidget.setItem(self.row, self.col, item)

        return super(InputElementTable, self).eventFilter(source, event)


# The window containing 2 lists with provision to add and subtract ions to decompose
# Inherits from Ui_DecomposeList
class InputDecomposeList(Ui_DecomposeList.Ui_Dialog, QDialog):
    def __init__(self, df_el=[], parent=None):
        super(InputDecomposeList, self).__init__(parent)
        self.setupUi(self)
        self.df_el = df_el
        self.df_decompose_el = None
        self.init_list()
        self.pushButton.clicked.connect(self.add_ion)
        self.pushButton_2.clicked.connect(self.subtract_ion)
        self.buttonBox.accepted.connect(self.submit_close)

    def submit_close(self):
        self.df_decompose_el = [str(self.listWidget_2.item(i).text()) for i in range(self.listWidget_2.count())]

    def init_list(self):
        for text in self.df_el:
            self.listWidget.addItem(text)

    def add_ion(self):
        if self.listWidget.currentItem():
            selected_text = self.listWidget.currentItem().text()
            self.listWidget_2.addItem(selected_text)
            self.listWidget.takeItem(self.listWidget.currentRow())

    def subtract_ion(self):
        if self.listWidget_2.currentItem():
            selected_text = self.listWidget_2.currentItem().text()
            self.listWidget.addItem(selected_text)
            self.listWidget_2.takeItem(self.listWidget_2.currentRow())


# The class inherits a custom UI layout and has facility to analyse mono-layers from H5 file
# Inherits from Ui_MonoLayer
class MonoLayerDialog(Ui_MonoLayer.Ui_Dialog, QDialog):
    def __init__(self, parent=None):
        super(MonoLayerDialog, self).__init__(parent)
        self.setupUi(self)

        self.hdf_file = None
        self.df_apt = None
        self.df_apt_layer = None
        self.widget_window = None
        self.ION = None
        self.my_old_plane = None
        self.alpha = None
        self.scatter3d = None
        self.scatter3d_noise_free = None
        self.scatter3d_noise = None
        self.count_layer = None
        self.dist_layer = None
        self.layer_thick_dict = None
        self.my_final_layers = None
        self.df_el = None
        self.df_decompose_el = None
        self.no_layers = False
        self.decomposition_list = None

        self.PeriodicTableCustom = PeriodicTableCustom()

        self.xs = 0
        self.ys = 0
        self.zs = 0
        self.d = 0
        self.x_plane_start = 0
        self.x_plane_end = 1
        self.y_plane_start = 0
        self.y_plane_end = 1
        self.z_plane_start = 0
        self.z_plane_end = 1
        self.plane = [0, 0, 0]

        # Progress bar
        self.completed = 0
        self.progressBar.setValue(self.completed)

        # Buttons
        self.pushButton_2.clicked.connect(self.plot_3d, False)  # Plot the APT data in 3D
        self.pushButton_3.clicked.connect(self.input_file)  # Input H5 file
        self.pushButton_4.clicked.connect(self.plot_plane)  # Plot the plane based on given miller indices (and d)
        self.pushButton_11.clicked.connect(self.plot_DBScan)  # Scatter plot the ions after doing DBScan
        self.pushButton_5.clicked.connect(self.find_layers)  # Start finding layers by travelling along plane direction
        self.pushButton.clicked.connect(self.show_peaks)  # Show the graph with ion counts across layer bins
        self.pushButton_13.clicked.connect(self.plot_layers)  # calculate the layer ions and plot layers
        self.pushButton_14.clicked.connect(self.decompose_list)  # give ions to decompose
        self.pushButton_16.clicked.connect(self.export_report)  # calculate and export final report as word docx
        self.pushButton_15.clicked.connect(self.export_hdf)  # export final dataframe as HDF file

        self.textEdit.setReadOnly(True)
        self.textEdit.mouseDoubleClickEvent = self.textEdit_click
        self.widget.fig = Figure()
        self.widget.canvas = FigureCanvas(self.widget.fig)
        self.widget.axes = self.widget.fig.add_subplot(111, projection='3d')

    # The below function is used to read the H5 file containing binned (mapped) apt data for mono-layer analysis
    def input_file(self):
        file = QFileDialog.getOpenFileName(self, 'Select HDF File"',
                                           os.getcwd(), "HDF files (*.h5)")
        try:
            self.hdf_file = file[0]
            self.df_apt = pd.read_hdf(self.hdf_file)
            self.pushButton_2.setEnabled(True)
            self.pushButton.setEnabled(True)

        except:
            self.pushButton_2.setEnabled(False)
            self.pushButton.setEnabled(False)

    # The below function is invoked upon doubleclick on the edit next to layer_element. It takes event as an input
    # The function is used to find the layer ion from periodic table to plot in 3D
    def textEdit_click(self, event):
        self.PeriodicTableCustom.show()
        if self.PeriodicTableCustom.exec() == 1:
            element_dict['charge'] = []
            element_dict['charge'].append(self.PeriodicTableCustom.emit_charge())
            text = ''
            for i in range(len(element_dict['ion'])):
                text = text + str(element_dict['ion'][i]) + common.subscript(str(element_dict['num'][i]))
            if text:
                charge = str(element_dict['charge'][0])
                charge = '(' + charge + ')'
                text = text + charge

            self.textEdit.setText(text)

    # Function to remove all plots after checking if they exist
    def remove_plots(self):
        if self.scatter3d:
            self.scatter3d.remove()
        del self.scatter3d
        self.scatter3d = None
        if self.scatter3d_noise_free:
            self.scatter3d_noise_free.remove()
        del self.scatter3d_noise_free
        self.scatter3d_noise_free = None
        if self.scatter3d_noise:
            self.scatter3d_noise.remove()
        del self.scatter3d_noise
        self.scatter3d_noise = None

    # The below function is used to plot the layer element in 3D plot. It can also plots after DBScan (keyword: outlier)
    # Initialises df_apt_layer which is required for further layer analysis, so must be always plotted
    def plot_3d(self, outliers):
        if self.df_apt is not None:
            layout = QVBoxLayout(self.widget)
            layout.addWidget(self.widget.canvas)

            # element_dict = {'ion': ['O', 'D'], 'num': ['1', '1'], 'mass': ['16.0', '2.014'], 'charge': ['1-']}

            if outliers:
                self.remove_plots()
                df_apt_layer_noise_free = self.df_apt_layer[self.df_apt_layer['DBSCAN_Label'] != -1]
                df_apt_layer_noise = self.df_apt_layer[self.df_apt_layer['DBSCAN_Label'] == -1]

                xs_nf = df_apt_layer_noise_free['X'].astype(float)
                ys_nf = df_apt_layer_noise_free['Y'].astype(float)
                zs_nf = df_apt_layer_noise_free['Z'].astype(float)

                xs_n = df_apt_layer_noise['X'].astype(float)
                ys_n = df_apt_layer_noise['Y'].astype(float)
                zs_n = df_apt_layer_noise['Z'].astype(float)

                self.scatter3d_noise_free = self.widget.axes.scatter(xs_nf, ys_nf, zs_nf, c='blue', s=30,
                                                                     label=self.ION, marker='.')

                self.scatter3d_noise = self.widget.axes.scatter(xs_n, ys_n, zs_n, c='red', s=30,
                                                                label=self.ION + "_noise", marker='.')

                self.widget.axes.set_xlim3d(self.x_plane_start, self.x_plane_end)
                self.widget.axes.set_ylim3d(self.y_plane_start, self.y_plane_end)
                self.widget.axes.set_zlim3d(self.z_plane_start, self.z_plane_end)

            else:
                self.remove_plots()

                if len(element_dict['ion']) > 0:
                    ion_array = element_dict['ion']
                    self.ION = ''
                    for ii in range(len(ion_array)):
                        self.ION = self.ION + ion_array[ii] + common.subscript(element_dict['num'][ii])
                        if ii == len(ion_array) - 1:
                            self.ION = self.ION + '(' + element_dict['charge'][0] + ')'

                    self.df_apt_layer = self.df_apt[self.df_apt['ION'] == self.ION]
                    # approximate input data as enclosed in 1st quadrant
                    # bringing negative coordinates into positive coordinates

                    if self.df_apt_layer.shape[0] > 2:
                        self.df_apt_layer = common.bring_df_to_positive_coord(self.df_apt_layer)

                        self.x_plane_start = self.df_apt_layer["X"].min()
                        self.x_plane_end = self.df_apt_layer["X"].max()
                        self.y_plane_start = self.df_apt_layer["Y"].min()
                        self.y_plane_end = self.df_apt_layer["Y"].max()
                        self.z_plane_start = self.df_apt_layer["Z"].min()
                        self.z_plane_end = self.df_apt_layer["Z"].max()

                        self.xs = self.df_apt_layer['X'].astype(float)
                        self.ys = self.df_apt_layer['Y'].astype(float)
                        self.zs = self.df_apt_layer['Z'].astype(float)

                    else:
                        self.xs = 0
                        self.ys = 0
                        self.zs = 0
                else:
                    self.xs = 0
                    self.ys = 0
                    self.zs = 0
                    self.ION = None

                self.scatter3d = self.widget.axes.scatter(self.xs, self.ys, self.zs, c='blue',
                                                          s=30, label=self.ION, marker='.')

                self.widget.axes.set_xlim3d(self.x_plane_start, self.x_plane_end)
                self.widget.axes.set_ylim3d(self.y_plane_start, self.y_plane_end)
                self.widget.axes.set_zlim3d(self.z_plane_start, self.z_plane_end)
                self.widget.axes.legend()
                self.widget.fig.canvas.draw()

            self.widget.axes.set_xlabel('X Axis')
            self.widget.axes.set_ylabel('Y Axis')
            self.widget.axes.set_zlabel('Z Axis')
            self.widget.fig.tight_layout()

    # The below function is used to get the plane coordinates in miller indices as well as plot it on the 3D layer plot
    # The default value for plane direction is set as [001] (parallel to Z axis)
    def plot_plane(self):
        self.plane = self.lineEdit.text()
        self.d = self.lineEdit_2.text()
        pattern_MI = "^['['](\s*[0-9]*?),(\s*[0-9]*?),(\s*[0-9]*)\]$"
        rex_pattern_MI = re.compile(pattern_MI)
        pattern_D = "^\s*\d+\.?\d*\s*$"
        rex_pattern_D = re.compile(pattern_D)

        if self.my_old_plane:
            self.my_old_plane.remove()
            self.my_old_plane = None

        if not self.d:
            self.d = 1
            self.lineEdit_2.setText("1")
        else:
            text = self.lineEdit_2.text()
            if rex_pattern_D.match(text):
                self.d = float(text)
            else:
                common.show_message("A (positive) float value is required as input for d intercept of the Layer Plane")
                self.d = None

        if not self.plane:
            self.plane = [0, 0, 1]
            self.lineEdit.setText("[0,0,1]")
        else:
            text = self.lineEdit.text()
            if rex_pattern_MI.match(text):
                text_groups = re.search(pattern_MI, text)
                a = float(text_groups.group(1))
                b = float(text_groups.group(2))
                c = float(text_groups.group(3))
                self.plane = [a, b, c]
            else:
                self.plane = None
                common.show_message("Incorrect format of Miller Indices input expected for Layer Plane (\"[a,b,c]\")")

        self.alpha = dict()  # alpha plane which curr to the plane specified by miller indices (inside unit cube)
        intercepts = ['a', 'b', 'c']

        if self.plane and self.d:
            for i in range(3):
                if self.plane[i] == 0:
                    self.alpha[intercepts[i]] = 0  # to avoid infinity
                else:
                    self.alpha[intercepts[i]] = 1 / self.plane[i]  # intercepts are reciprocals of miller indices
            self.alpha['d'] = self.d  # for ideal miller indices, d = 1

            X, Y, Z = None, None, None

            if self.alpha['c'] != 0:
                x = np.linspace(self.x_plane_start, self.x_plane_end, 10)
                y = np.linspace(self.y_plane_start, self.y_plane_end, 10)
                X, Y = np.meshgrid(x, y)
                Z = (self.alpha['d'] - self.alpha['a'] * X - self.alpha['b'] * Y) / self.alpha['c']

            else:
                if self.alpha['b'] != 0:
                    x = np.linspace(self.x_plane_start, self.x_plane_end, 10)
                    z = np.linspace(self.z_plane_start, self.z_plane_end, 10)
                    X, Z = np.meshgrid(x, z)
                    Y = (self.alpha['d'] - self.alpha['c'] * Z - self.alpha['a'] * X) / self.alpha['b']

                elif self.alpha['a'] != 0:
                    y = np.linspace(self.y_plane_start, self.y_plane_end, 10)
                    z = np.linspace(self.z_plane_start, self.z_plane_end, 10)
                    Y, Z = np.meshgrid(y, z)
                    X = (self.alpha['d'] - self.alpha['c'] * Z - self.alpha['b'] * Y) / self.alpha['a']

            if self.widget.axes:
                if self.my_final_layers:
                    for layer in self.my_final_layers:
                        layer.remove()
                    del self.my_final_layers
                    self.my_final_layers = []

                if X is not None:
                    my_plane = self.widget.axes.plot_surface(X, Y, Z, color='gray', shade=False)
                    self.my_old_plane = my_plane

                self.widget.fig.canvas.draw()

    # The below function optionally reduces the noise in the layer element using DBScan. Takes in Min_Points and Epsilon
    # It then plots the 3D plot with noises (ignored ions) as red dots. This helps in brute force optimisation of DBScan
    def plot_DBScan(self):
        if self.df_apt_layer is not None:
            pattern_Dbscan_minpoints = "^[0-9]*[1-9][0-9]*$"
            pattern_Dbscan_eps = "^\s*\d+\.?\d*\s*$"
            rex_pattern_Dbscan_minpoints = re.compile(pattern_Dbscan_minpoints)
            rex_pattern_Dbscan_eps = re.compile(pattern_Dbscan_eps)

            MIN_POINTS = None
            EPSILON = None

            min_pints_text = self.lineEdit_7.text()
            epsilon_text = self.lineEdit_8.text()

            if rex_pattern_Dbscan_minpoints.match(min_pints_text):
                MIN_POINTS = int(min_pints_text)
            else:
                common.show_message("a positive integer value is expected for MIN_POINTS")

            if rex_pattern_Dbscan_eps.match(epsilon_text):
                EPSILON = float(epsilon_text)
            else:
                common.show_message("a positive float value is expected for EPSILON")

            if type(MIN_POINTS) == int and MIN_POINTS > 0 and type(EPSILON) == float and EPSILON > 0:
                self.remove_plots()
                X = self.df_apt_layer[['X', 'Y', 'Z']].to_numpy()
                db = DBSCAN(eps=EPSILON, min_samples=MIN_POINTS).fit(X)
                labels = db.labels_
                self.df_apt_layer['DBSCAN_Label'] = labels.reshape(-1, 1)

                self.plot_3d(True)
                if self.scatter3d_noise is not None:
                    self.widget.fig.canvas.draw()

            if type(MIN_POINTS) == int and MIN_POINTS == 0 and type(EPSILON) == float and EPSILON == 0:
                self.remove_plots()
                if 'DBSCAN_Label' in self.df_apt_layer.columns:
                    self.df_apt_layer = self.df_apt_layer.drop('DBSCAN_Label', 1)
                self.plot_3d(False)
                if self.scatter3d_noise is not None:
                    self.widget.fig.canvas.draw()

        else:
            common.show_message("Plot the layer elements once before attempting to remove outliers")

    # The below function calculates the layers using input bin size and plane direction.
    # The output of the function are inputs for histogram 'peak' plot: self.dist_layer and self.count_layer
    def find_layers(self):
        self.completed = 0
        self.progressBar.setValue(self.completed)

        pattern_no_layers = "^[0-9]*[1-9][0-9]*$"
        rex_pattern_bin = re.compile(pattern_no_layers)
        text = self.lineEdit_3.text()
        self.no_layers = False
        if text == '':
            self.no_layers = 1
            self.lineEdit_3.setText('1')
        else:
            if rex_pattern_bin.match(text):
                self.no_layers = int(text)
            else:
                common.show_message("A (positive) integer is required as input for No: of Layer Plane(s)")
        if self.no_layers:
            self.lineEdit_3.setText(str(self.no_layers))

        text = self.lineEdit_9.text()
        pattern_bin = "^\s*\d+\.?\d*\s*$"
        rex_pattern_bin = re.compile(pattern_bin)
        layer_bin = False
        if text == '':
            layer_bin = 0.1
            self.lineEdit_9.setText("0.1")
        else:
            if rex_pattern_bin.match(text):
                layer_bin = float(text)
            else:
                common.show_message("A (positive) float value is required as input for bin of the Layer Plane")

        if layer_bin and self.alpha is not None and self.df_apt_layer is not None:
            if 'DBSCAN_Label' in self.df_apt_layer.columns:
                df_apt_layer_noise_free = self.df_apt_layer[self.df_apt_layer['DBSCAN_Label'] != -1]
            else:
                df_apt_layer_noise_free = self.df_apt_layer
            max_xyz = int(
                max(self.df_apt_layer['X'].max(), self.df_apt_layer['Y'].max(), self.df_apt_layer['Z'].max()) + 1)
            bin_num = int(max_xyz / layer_bin + 1)
            beta = [dict() for x in range(bin_num)]  # beta planes are number of binning planes
            self.count_layer = np.empty(bin_num)

            for i in range(bin_num):
                beta[i] = {k: self.alpha[k] for k in set(list(self.alpha.keys())) - {'d'}}
                beta[i]['d'] = (i + 1) * layer_bin

            # find number of points bw origin and beta[0]
            for bin_i in range(len(self.count_layer)):
                self.progressBar.setValue(self.completed)
                self.completed = (bin_i / len(self.count_layer)) * 100.0

                if bin_i == 0:
                    counts = 0
                    for i in range(df_apt_layer_noise_free.shape[0]):
                        p_xyz = [df_apt_layer_noise_free.iloc[i]['X'], df_apt_layer_noise_free.iloc[i]['Y'],
                                 df_apt_layer_noise_free.iloc[i]['Z']]
                        side = beta[0]['a'] * p_xyz[0] + beta[0]['b'] * p_xyz[1] + beta[0]['c'] * p_xyz[2] - beta[0][
                            'd']
                        side_sign = np.sign(side)
                        if side_sign == -1:
                            counts = counts + 1
                    self.count_layer[bin_i] = counts
                else:
                    counts = 0
                    for i in range(df_apt_layer_noise_free.shape[0]):
                        p_xyz = [df_apt_layer_noise_free.iloc[i]['X'], df_apt_layer_noise_free.iloc[i]['Y'],
                                 df_apt_layer_noise_free.iloc[i]['Z']]
                        side_beta_1 = beta[bin_i - 1]['a'] * p_xyz[0] + beta[bin_i - 1]['b'] * p_xyz[1] + \
                                      beta[bin_i - 1]['c'] * p_xyz[2] - beta[bin_i - 1]['d']
                        side_beta_2 = beta[bin_i]['a'] * p_xyz[0] + beta[bin_i]['b'] * p_xyz[1] + \
                                      beta[bin_i]['c'] * p_xyz[2] - beta[bin_i]['d']
                        side_sign_beta_1 = np.sign(side_beta_1)
                        side_sign_beta_2 = np.sign(side_beta_2)
                        if side_sign_beta_2 == -1 and side_sign_beta_1 == 1:
                            counts = counts + 1

                    self.count_layer[bin_i] = counts

                if bin_i == len(self.count_layer) - 1:
                    self.completed = 100
                self.progressBar.setValue(self.completed)

            self.dist_layer = np.linspace(0, max_xyz, num=bin_num)

    # The below function is invoked when we view the self.dist_layer vs self.count_layer peaks
    # This helps to make sure if there are in-fact mono layers present
    def show_peaks(self):
        if self.count_layer is not None:
            plt.figure(figsize=(10, 10), dpi=80, facecolor='w', edgecolor='k')
            yhat = savgol_filter(self.count_layer, 21, 3)
            plt.plot(self.dist_layer, self.count_layer, color='blue')
            plt.plot(self.dist_layer, yhat, color='red')
            # plt.plot(self.dist_layer, smooth(self.count_layer,2), 'green', lw=2)
            plt.show()

    # The below function finds and plots layers in 3D plot based on the conditions and number of layers
    def plot_layers(self):

        if self.no_layers and self.count_layer is not None:
            if self.widget.axes:
                if self.my_final_layers:
                    for layer in self.my_final_layers:
                        layer.remove()
                    del self.my_final_layers
                    self.my_final_layers = []
                    self.widget.fig.canvas.draw()

            count_index = self.count_layer.argsort()[-self.no_layers:][::-1]
            count_index = np.sort(count_index)
            self.layer_thick_dict = {'start': [], 'end': []}

            for i in range(len(count_index)):
                scan_iter = count_index[i]
                never_zero_flag = True
                while scan_iter != 0:
                    if self.count_layer[scan_iter] == 0:
                        never_zero_flag = False
                        self.layer_thick_dict['start'].append(self.dist_layer[scan_iter])
                        break
                    scan_iter = scan_iter - 1

                if never_zero_flag is True:
                    self.layer_thick_dict['start'].append(self.dist_layer[0])

                scan_iter = count_index[i]
                never_zero_flag = True
                while scan_iter != len(self.count_layer):
                    if self.count_layer[scan_iter] == 0:
                        never_zero_flag = False
                        self.layer_thick_dict['end'].append(self.dist_layer[scan_iter])
                        break
                    scan_iter = scan_iter + 1

                if never_zero_flag is True:
                    self.layer_thick_dict['end'].append(self.dist_layer[len(self.count_layer)])

            beta = [dict() for x in range(self.no_layers)]
            for ii in range(self.no_layers):
                beta[ii] = {k: self.alpha[k] for k in set(list(self.alpha.keys())) - {'d'}}
                beta[ii]['d'] = 0

            def check_beta_thick_high(row, bth, blh):
                side_high = np.sign(bth['a'] * row['X'] + bth['b'] * row['Y'] + bth['c'] * row['Z'] - bth['d'])
                side_low = np.sign(blh['a'] * row['X'] + blh['b'] * row['Y'] + blh['c'] * row['Z'] - blh['d'])

                side_status = False
                if side_high == -1 and side_low == 1:
                    side_status = True

                return side_status

            self.df_apt["layer_thick_start"] = np.nan
            self.df_apt["layer_thick_end"] = np.nan
            for ii in range(self.no_layers):
                beta_thick_high = {k: beta[ii][k] for k in set(list(beta[ii].keys())) - {'d'}}
                beta_thick_low = {k: beta[ii][k] for k in set(list(beta[ii].keys())) - {'d'}}
                beta_thick_low['d'] = self.layer_thick_dict['start'][ii]
                beta_thick_high['d'] = self.layer_thick_dict['end'][ii]

                if beta_thick_low['d'] < 0:
                    beta_thick_low['d'] = 0

                col_name = 'layer_' + str(ii + 1)
                self.df_apt[col_name] = self.df_apt.apply(check_beta_thick_high, bth=beta_thick_high,
                                                          blh=beta_thick_low,
                                                          axis=1)
                self.df_apt['layer_thick_start'] = np.where(self.df_apt[col_name] == True,
                                                            self.layer_thick_dict['start'][ii],
                                                            self.df_apt['layer_thick_start'])
                self.df_apt['layer_thick_end'] = np.where(self.df_apt[col_name] == True,
                                                          self.layer_thick_dict['end'][ii],
                                                          self.df_apt['layer_thick_end'])

            self.my_final_layers = []
            for ii in range(self.no_layers):
                h1 = self.layer_thick_dict['start'][ii]
                h2 = self.layer_thick_dict['end'][ii]

                X, Y, Z = None, None, None

                if self.alpha['c'] != 0:
                    x = np.linspace(self.x_plane_start, self.x_plane_end, 10)
                    y = np.linspace(self.y_plane_start, self.y_plane_end, 10)
                    X, Y = np.meshgrid(x, y)
                    Z_low = (h1 - self.alpha['a'] * X - self.alpha['b'] * Y) / self.alpha['c']
                    Z_high = (h2 - self.alpha['a'] * X - self.alpha['b'] * Y) / self.alpha['c']

                    if self.widget.axes:
                        if self.my_old_plane:
                            self.my_old_plane.remove()
                            self.my_old_plane = None

                        self.my_final_layers.append(
                            self.widget.axes.plot_surface(X, Y, Z_low, color='g', rstride=1, cstride=1, alpha=0.1))
                        self.my_final_layers.append(
                            self.widget.axes.plot_surface(X, Y, Z_high, color='g', rstride=1, cstride=1, alpha=0.1))
                        self.widget.fig.canvas.draw()

                else:
                    if self.alpha['b'] != 0:
                        x = np.linspace(self.x_plane_start, self.x_plane_end, 10)
                        z = np.linspace(self.z_plane_start, self.z_plane_end, 10)
                        X, Z = np.meshgrid(x, z)
                        Y_low = (h1 - self.alpha['c'] * Z - self.alpha['a'] * X) / self.alpha['b']
                        Y_high = (h2 - self.alpha['c'] * Z - self.alpha['a'] * X) / self.alpha['b']

                        if self.widget.axes:
                            if self.my_old_plane:
                                self.my_old_plane.remove()
                                self.my_old_plane = None

                            self.my_final_layers.append(
                                self.widget.axes.plot_surface(X, Y_low, Z, color='g', rstride=1, cstride=1, alpha=0.1))
                            self.my_final_layers.append(
                                self.widget.axes.plot_surface(X, Y_high, Z, color='g', rstride=1, cstride=1, alpha=0.1))
                            self.widget.fig.canvas.draw()


                    elif self.alpha['a'] != 0:
                        y = np.linspace(self.y_plane_start, self.y_plane_end, 10)
                        z = np.linspace(self.z_plane_start, self.z_plane_end, 10)
                        Y, Z = np.meshgrid(y, z)
                        X_low = (h1 - self.alpha['c'] * Z - self.alpha['b'] * Y) / self.alpha['a']
                        X_high = (h2 - self.alpha['c'] * Z - self.alpha['b'] * Y) / self.alpha['a']

                        if self.widget.axes:
                            if self.my_old_plane:
                                self.my_old_plane.remove()
                                self.my_old_plane = None

                            self.my_final_layers.append(
                                self.widget.axes.plot_surface(X_low, Y, Z, color='g', rstride=1, cstride=1, alpha=0.1))
                            self.my_final_layers.append(
                                self.widget.axes.plot_surface(X_high, Y, Z, color='g', rstride=1, cstride=1, alpha=0.1))
                            self.widget.fig.canvas.draw()

    # The following function is for optional entry where few input ions maybe specified to be decomposed in the report
    def decompose_list(self):
        if self.df_apt is not None:
            self.df_el = self.df_apt['ION'].unique()
            self.decomposition_list = InputDecomposeList(self.df_el)
            if self.decomposition_list.exec_():
                if self.decomposition_list.df_decompose_el:
                    self.df_decompose_el = self.decomposition_list.df_decompose_el

    # The final report containing counts of ions in the layers (with/without decomposition) is generated as a doc file
    def export_report(self):
        if self.no_layers and self.count_layer is not None:
            file = QFileDialog.getSaveFileName(self, 'Select/Create Word File"',
                                               os.getcwd(), "MS Word files (*.docx)")

            if len(file[0]) > 1:
                mydoc = docx.Document()
                self.widget.fig.savefig("temp.jpg")
                run = mydoc.add_paragraph().add_run()
                run.add_picture("temp.jpg")
                os.remove("temp.jpg")

                for i in range(self.no_layers):
                    df_layer = self.df_apt[self.df_apt['layer_' + str(i + 1)] == True]
                    df_ions = df_layer['ION'].value_counts().to_frame().reset_index()
                    layer_t1 = round(df_layer.iloc[0]['layer_thick_start'], 2)
                    layer_t2 = round(df_layer.iloc[0]['layer_thick_end'], 2)
                    df_ions = df_ions.rename(columns={"index": "ION", "ION": "counts"})

                    mydoc.add_paragraph("Layer %i: ion counts between %f nm and %f nm" % (i + 1, layer_t1, layer_t2))
                    mydoc.add_paragraph(df_ions.to_string())
                    dict_decomposed = {}
                    if self.df_decompose_el is not None:
                        mydoc.add_paragraph("The input Decompose ions are: %s" % str(self.df_decompose_el))
                        for ion in range(len(self.df_decompose_el)):
                            ion_text = self.df_decompose_el[ion]
                            if ion_text in self.df_apt['ION'].values:
                                df1 = df_layer[(df_layer['ION'].isin([ion_text]))]
                                if df1.shape[0] > 0:
                                    atom_count = len(df1.iloc[0]['ion'])
                                    for num in range(atom_count):
                                        ion_decomposed = df1.iloc[0]['ion'][num]
                                        count = int(df1.iloc[0]['num'][num]) * df1.shape[0]

                                        if str(ion_decomposed) in dict_decomposed:
                                            dict_decomposed[str(ion_decomposed)] = int(
                                                dict_decomposed[str(ion_decomposed)]) + int(count)
                                        else:
                                            dict_decomposed[str(ion_decomposed)] = int(count)

                        for sl_no, dict_ion in enumerate(dict_decomposed):
                            mydoc.add_paragraph("%i) Sum of %s = %i" % (sl_no + 1, dict_ion, dict_decomposed[dict_ion]))
                mydoc.save(file[0])

        else:
            common.show_message("No valid report to output into the .docx file")

    # The final dataframe after layer thickness and layer category (boolean) classification is output as an hdf file
    def export_hdf(self):
        if self.no_layers and self.count_layer is not None:
            file = QFileDialog.getSaveFileName(self, 'Select/Create HDF File"', os.getcwd(), "HDF files (*.h5)")
            if len(file[0]) > 1:
                float_columns = ['X', 'Y', 'Z', 'MN_Ratio', 'peak_MNRatio', 'peak_max_cutoff_width',
                                 'layer_thick_start', 'layer_thick_end']
                int_columns = ['peak_no', 'XYZ_total_count']
                bool_columns = []
                for i in range(1, self.no_layers + 1):
                    bool_columns.append('layer_' + str(i))

                self.df_apt.loc[:, float_columns] = self.df_apt[float_columns].applymap(float)
                self.df_apt.loc[:, int_columns] = self.df_apt[int_columns].applymap(int)
                self.df_apt.loc[:, bool_columns] = self.df_apt[bool_columns].applymap(bool)

                self.df_apt.to_hdf(file[0], key='df_apt_final', mode='w')

        else:
            common.show_message("No valid data to output into the .h5 file")


# The class inherits a custom UI layout and has facility to analyse abstract-layers from H5 file
# Inherits from Ui_AbstractLayer
class AbstractLayerDialog(Ui_AbstractLayer.Ui_Dialog, QDialog):
    def __init__(self, parent=None):
        super(AbstractLayerDialog, self).__init__(parent)
        self.setupUi(self)

        self.hdf_file = None
        self.df_apt = None
        self.scatter3d = None
        self.scatter3d_noise_free = None
        self.scatter3d_noise = None
        self.Hull3d_lines = None
        self.decomposition_list = None
        self.df_apt_layer = None
        self.df_el = None
        self.df_apt_final = None
        self.df_decompose_el = None
        self.ION = None

        self.x_plane_start = 0
        self.x_plane_end = 1
        self.y_plane_start = 0
        self.y_plane_end = 1
        self.z_plane_start = 0
        self.z_plane_end = 1

        self.xs = 0
        self.ys = 0
        self.zs = 0

        self.PeriodicTableCustom = PeriodicTableCustom()
        self.textEdit.setReadOnly(True)
        self.textEdit.mouseDoubleClickEvent = self.textEdit_click
        self.widget.fig = Figure()
        self.widget.canvas = FigureCanvas(self.widget.fig)
        self.widget.axes = self.widget.fig.add_subplot(111, projection='3d')

        self.pushButton_3.clicked.connect(self.input_file)  # Input H5 file
        self.pushButton_2.clicked.connect(self.plot_3d, False)  # Plot the APT data in 3D
        self.pushButton_11.clicked.connect(self.plot_DBScan)  # Scatter plot the ions after doing DBScan
        self.pushButton_5.clicked.connect(self.plot_ConvexHull)  # Scatter plot the ions after doing DBScan
        self.pushButton_14.clicked.connect(self.decompose_list)  # give ions to decompose
        self.pushButton_16.clicked.connect(self.export_report)  # calculate and export final report as word docx
        self.pushButton_15.clicked.connect(self.export_hdf)  # export final dataframe as HDF file

    # The below function is used to read the H5 file containing binned (mapped) apt data for abstract-layer analysis
    def input_file(self):
        # file = ['C://Users/arjun/Downloads/APT_Code/APT_Project/totaldata3_binned.h5']
        file = QFileDialog.getOpenFileName(self, 'Select HDF File"',
                                           os.getcwd(), "HDF files (*.h5)")
        try:
            self.hdf_file = file[0]
            self.df_apt = pd.read_hdf(self.hdf_file)
            self.pushButton_2.setEnabled(True)

        except:
            self.pushButton_2.setEnabled(False)

    # The below function is invoked upon doubleclick on the edit next to layer_element. It takes event as an input
    # The function is used to find the layer ion from periodic table to plot in 3D
    def textEdit_click(self, event):
        self.PeriodicTableCustom.show()
        if self.PeriodicTableCustom.exec() == 1:
            element_dict['charge'] = []
            element_dict['charge'].append(self.PeriodicTableCustom.emit_charge())
            text = ''
            for i in range(len(element_dict['ion'])):
                text = text + str(element_dict['ion'][i]) + common.subscript(str(element_dict['num'][i]))
            if text:
                charge = str(element_dict['charge'][0])
                charge = '(' + charge + ')'
                text = text + charge

            self.textEdit.setText(text)

    # Function to remove all plots after checking if they exist
    def remove_plots(self):
        if self.scatter3d:
            self.scatter3d.remove()
        del self.scatter3d
        self.scatter3d = None
        if self.scatter3d_noise_free:
            self.scatter3d_noise_free.remove()
        del self.scatter3d_noise_free
        self.scatter3d_noise_free = None
        if self.scatter3d_noise:
            self.scatter3d_noise.remove()
        del self.scatter3d_noise
        self.scatter3d_noise = None

    # The below function is used to plot the layer element in 3D plot. It can also plots after DBScan (keyword: outlier)
    # Initialises df_apt_layer which is required for further layer analysis, so must be always plotted
    def plot_3d(self, outliers):
        if self.df_apt is not None:
            layout = QVBoxLayout(self.widget)
            layout.addWidget(self.widget.canvas)

            # element_dict = {'ion': ['D', 'O'], 'num': ['2', '1'], 'mass': ['2.014', '16.0'], 'charge': ['1+']}

            if outliers:
                self.remove_plots()
                df_apt_layer_noise_free = self.df_apt_layer[self.df_apt_layer['DBSCAN_Label'] != -1]
                df_apt_layer_noise = self.df_apt_layer[self.df_apt_layer['DBSCAN_Label'] == -1]

                xs_nf = df_apt_layer_noise_free['X'].astype(float)
                ys_nf = df_apt_layer_noise_free['Y'].astype(float)
                zs_nf = df_apt_layer_noise_free['Z'].astype(float)

                xs_n = df_apt_layer_noise['X'].astype(float)
                ys_n = df_apt_layer_noise['Y'].astype(float)
                zs_n = df_apt_layer_noise['Z'].astype(float)

                self.scatter3d_noise_free = self.widget.axes.scatter(xs_nf, ys_nf, zs_nf, c='blue', s=30,
                                                                     label=self.ION, marker='.')

                self.scatter3d_noise = self.widget.axes.scatter(xs_n, ys_n, zs_n, c='red', s=30,
                                                                label=self.ION + "_noise", marker='.')

                self.widget.axes.set_xlim3d(self.x_plane_start, self.x_plane_end)
                self.widget.axes.set_ylim3d(self.y_plane_start, self.y_plane_end)
                self.widget.axes.set_zlim3d(self.z_plane_start, self.z_plane_end)

            else:
                self.remove_plots()
                if len(element_dict['ion']) > 0:
                    ion_array = element_dict['ion']
                    self.ION = ''
                    for ii in range(len(ion_array)):
                        self.ION = self.ION + ion_array[ii] + common.subscript(element_dict['num'][ii])
                        if ii == len(ion_array) - 1:
                            self.ION = self.ION + '(' + element_dict['charge'][0] + ')'

                    self.df_apt_layer = self.df_apt[self.df_apt['ION'] == self.ION]
                    # approximate input data as enclosed in 1st quadrant
                    # bringing negative coordinates into positive coordinates
                    if self.df_apt_layer.shape[0] > 2:
                        self.df_apt_layer = common.bring_df_to_positive_coord(self.df_apt_layer)

                        self.x_plane_start = self.df_apt_layer["X"].min()
                        self.x_plane_end = self.df_apt_layer["X"].max()
                        self.y_plane_start = self.df_apt_layer["Y"].min()
                        self.y_plane_end = self.df_apt_layer["Y"].max()
                        self.z_plane_start = self.df_apt_layer["Z"].min()
                        self.z_plane_end = self.df_apt_layer["Z"].max()

                        self.xs = self.df_apt_layer['X'].astype(float)
                        self.ys = self.df_apt_layer['Y'].astype(float)
                        self.zs = self.df_apt_layer['Z'].astype(float)
                    else:
                        self.xs = 0
                        self.ys = 0
                        self.zs = 0
                else:
                    self.xs = 0
                    self.ys = 0
                    self.zs = 0
                    self.ION = None

                self.scatter3d = self.widget.axes.scatter(self.xs, self.ys, self.zs, c='blue',
                                                          s=30, label=self.ION, marker='.')

                self.widget.fig.subplots_adjust(0.2, 0.2, 0.8, 0.8)  # left,bottom,right,top

                self.widget.axes.set_xlim3d(self.x_plane_start, self.x_plane_end)
                self.widget.axes.set_ylim3d(self.y_plane_start, self.y_plane_end)
                self.widget.axes.set_zlim3d(self.z_plane_start, self.z_plane_end)
                self.widget.axes.legend()
                self.widget.fig.canvas.draw()
                self.widget.axes.set_xlabel('X Axis')
                self.widget.axes.set_ylabel('Y Axis')
                self.widget.axes.set_zlabel('Z Axis')
                self.widget.fig.tight_layout()

    # The below function optionally reduces the noise in the layer element using DBScan. Takes in Min_Points and Epsilon
    # It then plots the 3D plot with noises (ignored ions) as red dots. This helps in brute force optimisation of DBScan
    def plot_DBScan(self):
        if self.df_apt_layer is not None:
            pattern_Dbscan = "^(0|[1-9]\d*)?(\.\d+)?(?<=\d)$"
            rex_pattern_Dbscan = re.compile(pattern_Dbscan)

            MIN_POINTS = None
            EPSILON = None

            # self.lineEdit_7.setText("3")
            # self.lineEdit_8.setText("0.75")

            min_pints_text = self.lineEdit_7.text()
            epsilon_text = self.lineEdit_8.text()

            if rex_pattern_Dbscan.match(min_pints_text):
                MIN_POINTS = int(min_pints_text)
            else:
                common.show_message("a positive integer value is expected for MIN_POINTS")

            if rex_pattern_Dbscan.match(epsilon_text):
                EPSILON = float(epsilon_text)
            else:
                common.show_message("a positive float value is expected for EPSILON")

            if type(MIN_POINTS) == int and MIN_POINTS > 0 and type(EPSILON) == float and EPSILON > 0:
                self.remove_plots()

                X = self.df_apt_layer[['X', 'Y', 'Z']].to_numpy()
                db = DBSCAN(eps=EPSILON, min_samples=MIN_POINTS).fit(X)
                labels = db.labels_
                self.df_apt_layer['DBSCAN_Label'] = labels.reshape(-1, 1)

                if self.radioButton.isChecked():
                    self.df_apt_layer.loc[self.df_apt_layer['DBSCAN_Label'] != -1, 'DBSCAN_Label'] = 2
                    self.df_apt_layer.loc[self.df_apt_layer['DBSCAN_Label'] == -1, 'DBSCAN_Label'] = 1
                    self.df_apt_layer.loc[self.df_apt_layer['DBSCAN_Label'] == 2, 'DBSCAN_Label'] = -1

                self.plot_3d(True)
                if self.scatter3d_noise is not None:
                    self.widget.fig.canvas.draw()

            if type(MIN_POINTS) == int and MIN_POINTS == 0 and type(EPSILON) == float and EPSILON == 0:
                self.remove_plots()

                if 'DBSCAN_Label' in self.df_apt_layer.columns:
                    self.df_apt_layer = self.df_apt_layer.drop('DBSCAN_Label', 1)
                self.plot_3d(False)
                if self.scatter3d_noise is not None:
                    self.widget.fig.canvas.draw()

        else:
            common.show_message("Plot the layer elements once before attempting to remove outliers")

    # The below function calculates the convex hull for the given layer element with/without DBScan
    # It also plots the updated convex hull as dotted green lines (takes time)
    def plot_ConvexHull(self):
        if self.df_apt_layer is not None:
            if self.widget.axes:
                if self.Hull3d_lines:
                    for layer in self.Hull3d_lines:
                        layer.remove()
                    del self.Hull3d_lines
                    self.Hull3d_lines = []
                    self.widget.fig.canvas.draw()

            if 'DBSCAN_Label' in self.df_apt_layer.columns:
                df_apt_layer_noise_free = self.df_apt_layer[self.df_apt_layer['DBSCAN_Label'] != -1]
            else:
                df_apt_layer_noise_free = self.df_apt_layer

            df_apt_layer_noise_free['Status_convex_hull'] = True
            points_not_noise = df_apt_layer_noise_free[['X', 'Y', 'Z']].values
            hull = ConvexHull(points_not_noise)
            deln = Delaunay(points_not_noise[hull.vertices])

            if self.df_apt is not None:
                self.df_apt = common.bring_df_to_positive_coord(self.df_apt)

            df_apt_non_layer = self.df_apt[~self.df_apt['ION'].isin([self.ION])]

            non_layer_points = df_apt_non_layer[['X', 'Y', 'Z']].to_numpy()
            df_apt_non_layer['Status_convex_hull'] = deln.find_simplex(non_layer_points) >= 0

            frames = [df_apt_layer_noise_free, df_apt_non_layer]
            self.df_apt_final = pd.concat(frames)
            self.Hull3d_lines = []
            for s in hull.simplices:
                x, y, z = points_not_noise[s, 0], points_not_noise[s, 1], points_not_noise[s, 2]
                self.Hull3d_lines.append(self.widget.axes.plot(x, y, z, '--', c='green', alpha=0.1)[0])
                self.widget.fig.canvas.draw()

    # The following function is for optional entry where few input ions maybe specified to be decomposed in the report
    def decompose_list(self):
        if self.df_apt is not None:
            self.df_el = self.df_apt['ION'].unique()
            self.decomposition_list = InputDecomposeList(self.df_el)
            if self.decomposition_list.exec_():
                if self.decomposition_list.df_decompose_el:
                    self.df_decompose_el = self.decomposition_list.df_decompose_el

    # The final report containing counts of ions in the layers (with/without decomposition) is generated as a doc file
    def export_report(self):
        if self.df_apt_final is not None:
            file = QFileDialog.getSaveFileName(self, 'Select/Create Word File"',
                                               os.getcwd(), "MS Word files (*.docx)")

            if len(file[0]) > 1:
                mydoc = docx.Document()
                if self.df_apt_final.shape[0] > 2:
                    df_layer = self.df_apt_final[self.df_apt_final['Status_convex_hull'] == True]
                    df_ions = df_layer['ION'].value_counts().to_frame().reset_index()
                    df_ions = df_ions.rename(columns={"index": "ION", "ION": "counts"})

                    self.widget.fig.savefig("temp.jpg")
                    run = mydoc.add_paragraph().add_run()
                    run.add_picture("temp.jpg")
                    os.remove("temp.jpg")

                    mydoc.add_paragraph(
                        "Calculating ion counts inside abstract layer formed by %s ion" % self.ION)

                    table = mydoc.add_table(rows=1, cols=df_ions.shape[1])

                    table_row = table.rows[0].cells
                    table_row[0].text = 'ION'
                    table_row[1].text = 'counts'

                    for _, row in df_ions.iterrows():
                        table_row = table.add_row().cells
                        table_row[0].text = str(row['ION'])
                        table_row[1].text = str(row['counts'])

                    table.style = 'Table Grid'

                    dict_decomposed = {}
                    if self.df_decompose_el is not None:
                        mydoc.add_paragraph("The input Decompose ions are: %s" % str(self.df_decompose_el))
                        for ion in range(len(self.df_decompose_el)):
                            ion_text = self.df_decompose_el[ion]
                            if ion_text in self.df_apt_final['ION'].values:
                                df1 = df_layer[(df_layer['ION'].isin([ion_text]))]
                                if df1.shape[0] > 0:
                                    atom_count = len(df1.iloc[0]['ion'])
                                    for num in range(atom_count):
                                        ion_decomposed = df1.iloc[0]['ion'][num]
                                        count = int(df1.iloc[0]['num'][num]) * df1.shape[0]

                                        if str(ion_decomposed) in dict_decomposed:
                                            dict_decomposed[str(ion_decomposed)] = int(
                                                dict_decomposed[str(ion_decomposed)]) + int(count)
                                        else:
                                            dict_decomposed[str(ion_decomposed)] = int(count)

                        for sl_no, dict_ion in enumerate(dict_decomposed):
                            mydoc.add_paragraph(
                                "%i) Sum of %s = %i" % (sl_no + 1, dict_ion, dict_decomposed[dict_ion]))
                mydoc.save(file[0])

        else:
            common.show_message("No valid report to output into the .docx file")

    # The final dataframe after abstract layer is detected is output as an hdf file
    def export_hdf(self):
        if self.df_apt_final is not None and self.df_apt_final.shape[0] > 0:
            file = QFileDialog.getSaveFileName(self, 'Select/Create HDF File"', os.getcwd(), "HDF files (*.h5)")
            if len(file[0]) > 1:
                float_columns = ['X', 'Y', 'Z', 'MN_Ratio', 'peak_MNRatio', 'peak_max_cutoff_width',
                                 'Status_convex_hull']
                int_columns = ['peak_no', 'XYZ_total_count']
                bool_columns = []
                for i in range(1, self.no_layers + 1):
                    bool_columns.append('layer_' + str(i))

                self.df_apt_final.loc[:, float_columns] = self.df_apt_final[float_columns].applymap(float)
                self.df_apt_final.loc[:, int_columns] = self.df_apt_final[int_columns].applymap(int)
                self.df_apt_final.loc[:, bool_columns] = self.df_apt_final[bool_columns].applymap(bool)

                self.df_apt_final.to_hdf(file[0], key='df_apt_final', mode='w')

        else:
            common.show_message("No valid data to output into the .h5 file")


# The class inherits a custom UI layout and has facility to analyse SRO parameter from H5 file
# Inherits from Ui_SRO
class SRODialog(Ui_SRO.Ui_Dialog, QDialog):
    def __init__(self, parent=None):
        super(SRODialog, self).__init__(parent)
        self.setupUi(self)

        self.df_apt = None
        self.poscar = None
        self.input_radius_table = None
        self.voxel_status = False
        self.r_star_array_count = []
        self.r_star_array = []
        self.gm_sro_array = []

        self.max_shell = 10  # maximum number of shells to be previewed and used
        self.num_accuracy_3dmf = 1000  # number of meshing points of plots for finding intersection in 3DMF calculation

        # Progress bar
        self.completed = 0
        self.progressBar.setValue(self.completed)

        self.pushButton_2.setEnabled(False)
        self.pushButton_5.setEnabled(False)
        self.pushButton_6.setEnabled(False)
        self.pushButton_7.setEnabled(False)
        self.pushButton_9.setEnabled(False)

        self.PeriodicTableCustom = PeriodicTableCustom()

        self.pushButton_1.clicked.connect(self.input_file)  # Input H5 file
        self.pushButton_4.clicked.connect(self.input_POSCAR)  # Input POSCAR file
        self.pushButton_5.clicked.connect(self.use_nearest_neighbour_chart)  # use the radius based on NN chart
        self.pushButton_6.clicked.connect(self.use_nearest_neighbour_3DMF)  # use the radius based on 3DMF chart
        self.pushButton_7.clicked.connect(self.use_manual_critical_radius)  # use the radius entered manually
        self.pushButton_2.clicked.connect(self.voxelize)  # Voxel the coordinates into cube of input dimension

        self.pushButton_8.clicked.connect(self.plot_nearest_neighbour_chart)  # plot NN chart
        self.pushButton_10.clicked.connect(self.plot_nearest_neighbour_3DMF)  # plot R* chart

        self.pushButton_11.clicked.connect(self.add_Bj)  # Add a Bj ion from periodic table
        self.pushButton_12.clicked.connect(self.Subtract_Bj)  # Subtract a Bj ion from periodic table
        self.pushButton_13.clicked.connect(self.add_Cl)  # Add a Cl ion from periodic table
        self.pushButton_14.clicked.connect(self.Subtract_Cl)  # Subtract a Cl ion from periodic table

        self.pushButton_9.clicked.connect(self.Calculate_SRO)  # Subtract a Cl ion from periodic table
        self.pushButton_15.clicked.connect(self.plot_GMSRO)  # To plot the final GM-SRO plot

    # The function used to read the H5 file containing binned (mapped) apt data
    def input_file(self):
        # file = ['C://Users/arjun/Downloads/APT_Code/APT_Project/totaldata2_binned.h5']
        file = QFileDialog.getOpenFileName(self, 'Select HDF File"',
                                           os.getcwd(), "HDF files (*.h5)")
        try:
            self.hdf_file = file[0]
            self.df_apt = pd.read_hdf(self.hdf_file)
            self.df_apt['status_Bj'] = None
            self.df_apt['status_Cl'] = None
            self.df_apt = common.bring_df_to_positive_coord(self.df_apt)
            self.df_apt['dist_origin_sq'] = pow(self.df_apt['X'], 2) + pow(self.df_apt['Y'], 2) + pow(
                self.df_apt['Z'], 2)
            self.df_apt['dist_origin'] = np.sqrt((self.df_apt['dist_origin_sq'].astype(float)))

            self.lineEdit.setText('2.0')
            self.lineEdit_2.setText('0.2')
            self.lineEdit_3.setText('3')

            self.pushButton_2.setEnabled(True)
            self.voxel_status = False

        except:
            self.pushButton_2.setEnabled(False)

    # The function used to read the POSCAR file for finding nearest neighbour distance and number
    def input_POSCAR(self):
        # file = ('C:/Users/arjun/Downloads/APT_Code/APT_Project/POSCAR.mp-153_Mg', 'POSCAR files (*.*)')
        file = QFileDialog.getOpenFileName(self, 'Select POSCAR File"',
                                           os.getcwd(), "POSCAR files (*.*)")

        try:
            atoms = read(file[0], format="vasp", )
            if atoms is not None:
                self.poscar_dist = file[0]
                self.pushButton_5.setEnabled(True)
                self.pushButton_6.setEnabled(True)
                self.textEdit_7.setText('0.015')
                self.textEdit_8.setText('0.015')

        except:
            common.show_message("Enter a valid POSCAR file")
            self.pushButton_5.setEnabled(False)
            self.pushButton_6.setEnabled(False)

    # The function input a critical radius and returns nearest distances & neighbour counts from existing POSCAR file
    def find_NearestNeighbours(self, cutoff_radius):
        if self.poscar_dist is not None:
            atoms = read(self.poscar_dist, format="vasp", )
            atoms.pbc = True
            nl = NeighborList([5] * len(atoms), primitive=NewPrimitiveNeighborList, self_interaction=False,
                              bothways=True)
            nl.update(atoms)
            dist = []

            for i in range(len(atoms)):
                indices, offsets = nl.get_neighbors(i)
                dist = []
                for j, offset in zip(indices, offsets):
                    distance_value = np.linalg.norm(
                        atoms.positions[j] + np.dot(offset, atoms.get_cell()) - atoms.get_positions()[i])
                    # The units are converted to Nm from A
                    dist.append(distance_value * 0.1)

            dist1 = np.around(dist, decimals=4)
            a = Counter(dist1)
            a1 = sorted(a.items())
            dist, count = zip(*a1)

            dist_new = []
            count_new = []
            dist_new.append(dist[0])
            count_new.append(count[0])

            for i in range(1, len(dist)):
                if (dist[i] - dist[i - 1]) < cutoff_radius:
                    count_new[-1] = count_new[-1] + count[i]
                else:
                    dist_new.append(dist[i])
                    count_new.append(count[i])

            return dist_new, count_new

    # Function to use the standard nearest neighbour chart with optional cutoff
    def use_nearest_neighbour_chart(self):
        pattern_cutoff_radius = "^(0|[1-9]\d*)?(\.\d+)?(?<=\d)$"
        rex_pattern_cutoff_radius = re.compile(pattern_cutoff_radius)
        cutoff_radius = self.textEdit_7.toPlainText()

        if rex_pattern_cutoff_radius.match(cutoff_radius):
            cutoff_radius = float(cutoff_radius)
        else:
            common.show_message("a positive float value is expected for cutoff radius")

        dist, count = self.find_NearestNeighbours(cutoff_radius)
        self.r_star_array = dist
        self.r_star_array_count = count
        self.pushButton_9.setEnabled(True)
        self.pushButton_7.setEnabled(True)

    # only plots the nearest neighbour chart based on POSCAR file
    def plot_nearest_neighbour_chart(self):
        pattern_cutoff_radius = "^(0|[1-9]\d*)?(\.\d+)?(?<=\d)$"
        rex_pattern_cutoff_radius = re.compile(pattern_cutoff_radius)
        cutoff_radius = self.textEdit_7.toPlainText()

        if rex_pattern_cutoff_radius.match(cutoff_radius):
            cutoff_radius = float(cutoff_radius)
        else:
            common.show_message("a positive float value is expected for cutoff radius")

        dist, count = self.find_NearestNeighbours(cutoff_radius)
        plt.title('Nearest Neighbour Chart')

        plt.stem(dist, count, linefmt='grey', markerfmt='o', use_line_collection=True)
        # plt.bar(dist, count)
        plt.show()

    # Function to use the 3DMF based nearest neighbour chart with optional cutoff
    def use_nearest_neighbour_3DMF(self):
        self.r_star_array = []
        self.r_star_array_count = []

        def find_rstar(dist, m, n1, n2, num, FWHM):
            r_star = np.nan
            if m > n2:
                return (np.nan)

            rmax_list = np.linspace(0, dist[n2 + 1], num)
            prob_m_list = [[0 for _ in range(num)] for _ in range(n2)]
            pm_list = []
            pm_n_list = []
            prob_n_complete = []

            def find_pm(rmax, m, dist):
                delta_res = FWHM / (2 * math.sqrt(2 * np.log(2)))

                dm = dist[m - 1]
                gm = abs((rmax - dm) / delta_res)
                z = gm / np.sqrt(2)

                if rmax < dm:
                    pm = (1 - special.erf(z)) / 2.0
                else:
                    pm = (1 + special.erf(z)) / 2.0

                return pm

            for i in range(0, n2):
                for j in range(len(rmax_list)):
                    rmax = rmax_list[j]
                    prob_m_list[i][j] = find_pm(rmax, i + 1, dist)
                    if i == m - 1:
                        pm_list.append(prob_m_list[i][j])

            for i in range(len(rmax_list)):
                rmax = rmax_list[i]
                deno = 0
                for j in range(n1, n2):
                    deno += count[j + 1] * find_pm(rmax, j + 1, dist)
                prob_n_complete.append(deno)

            for i in range(len(rmax_list)):
                p_num = count[m - 1] * pm_list[i]
                p_deno = prob_n_complete[i]
                pm_n_list.append(p_num / p_deno)

            idx = np.argwhere(np.diff(np.sign(np.asarray(pm_n_list) - np.asarray(pm_list)))).flatten()
            if len(idx) > 1:
                # print("intersecting points are: ", idx)
                idx = [idx[-1]]

            if not idx:
                r_star = np.nan

            if idx:
                r_star = rmax_list[idx[0]]

            return r_star

        pattern_cutoff_radius = "^(0|[1-9]\d*)?(\.\d+)?(?<=\d)$"
        rex_pattern_cutoff_radius = re.compile(pattern_cutoff_radius)
        cutoff_radius = self.textEdit_8.toPlainText()
        cutoff_radius_status = False

        if rex_pattern_cutoff_radius.match(cutoff_radius):
            cutoff_radius = float(cutoff_radius)
            cutoff_radius_status = True
        else:
            common.show_message("a positive float value is expected for cutoff radius")

        pattern_fwhm = "^(0|[1-9]\d*)?(\.\d+)?(?<=\d)$"
        rex_pattern_fwhm = re.compile(pattern_fwhm)
        FWHM = self.lineEdit_2.text()
        FWHM_status = False

        if rex_pattern_fwhm.match(FWHM):
            FWHM = float(FWHM)
            FWHM_status = True
        else:
            common.show_message("a positive float value is expected for FWHM")

        if FWHM_status and cutoff_radius_status:
            dist, count = self.find_NearestNeighbours(cutoff_radius)
            self.r_star_array_count = count
            for shell in range(1, self.max_shell + 5):
                self.r_star_array.append(find_rstar(dist, m=shell, n1=0, n2=10, num=self.num_accuracy_3dmf, FWHM=FWHM))

        self.pushButton_9.setEnabled(True)
        self.pushButton_7.setEnabled(True)

    # only plots the nearest neighbour chart based on 3DMF
    def plot_nearest_neighbour_3DMF(self):
        self.use_nearest_neighbour_3DMF()
        dist = self.r_star_array
        count = self.r_star_array_count[0:len(dist)]
        plt.title('Nearest Neighbour Chart')
        plt.stem(dist[0:self.max_shell], count[0:self.max_shell], linefmt='grey', markerfmt='o',
                 use_line_collection=True)
        plt.show()

    # Function to display chosen radius chart and make manual changes if needed
    def use_manual_critical_radius(self):
        self.input_radius_table = InputElementTable()
        self.input_radius_table.tableWidget.setColumnCount(2)
        self.input_radius_table.tableWidget.setHorizontalHeaderItem(0, QTableWidgetItem("Radius"))
        self.input_radius_table.tableWidget.setHorizontalHeaderItem(1, QTableWidgetItem("No: of Neighbours"))
        header = self.input_radius_table.tableWidget.horizontalHeader()
        header.setSectionResizeMode(1, QHeaderView.Stretch)
        header.setSectionResizeMode(0, QHeaderView.ResizeToContents)
        self.input_radius_table.pushButton_2.setVisible(False)
        self.input_radius_table.pushButton.setVisible(False)
        self.input_radius_table.tableWidget.cellClicked.disconnect()
        self.input_radius_table.tableWidget.cellDoubleClicked.disconnect()

        for row in range(self.max_shell):
            item_rad = QTableWidgetItem()
            item_count = QTableWidgetItem()
            item_count.setText(str(self.r_star_array_count[row]))
            self.input_radius_table.tableWidget.setItem(row, 1, item_count)
            item_rad.setText(str(self.r_star_array[row]))
            self.input_radius_table.tableWidget.setItem(row, 0, item_rad)
            self.input_radius_table.tableWidget.item(row, 0).setBackground(QtGui.QColor(255, 255, 255))
            if row > 0 and self.r_star_array[row - 1] > self.r_star_array[row]:
                self.input_radius_table.tableWidget.item(row, 0).setBackground(QtGui.QColor(150, 75, 75))

        self.input_radius_table.show()

        if self.input_radius_table.exec() == 1:
            r_star_array_temp = []
            for row in range(self.max_shell):
                item_rad = self.input_radius_table.tableWidget.item(row, 0)
                rad = float(item_rad.text())
                r_star_array_temp.append(rad)

            self.r_star_array = r_star_array_temp

    # Function to add a specified Bj from the list and show it on the line_edit widget
    def add_Bj(self):
        # text = "Gd₁(2+)"
        # self.df_apt.loc[self.df_apt['ION'] == text, 'status_Bj'] = True
        # self.lineEdit_4.setText(text + ",")

        if self.df_apt is not None:
            self.PeriodicTableCustom.show()
            if self.PeriodicTableCustom.exec() == 1:
                element_dict['charge'] = []
                element_dict['charge'].append(self.PeriodicTableCustom.emit_charge())
                text = ''
                for i in range(len(element_dict['ion'])):
                    text = text + str(element_dict['ion'][i]) + common.subscript(str(element_dict['num'][i]))
                if text:
                    charge = str(element_dict['charge'][0])
                    charge = '(' + charge + ')'
                    text = text + charge
                if self.df_apt[self.df_apt['ION'] == text].shape[0] > 0:
                    prev_text = self.lineEdit_4.text()
                    self.lineEdit_4.setText(prev_text + text + ",")
                    self.df_apt.loc[self.df_apt['ION'] == text, 'status_Bj'] = True
                else:
                    prev_text = self.lineEdit_4.text()
                    self.lineEdit_4.setText(prev_text + "N.A.,")
                    common.show_message("The entered element/ion does not exist in the input dataframe")

        else:
            common.show_message("Add the H5 data before addition of Bj/Cl ions")

    # Function to subtract a specified Bj from the list and show it on the line_edit widget
    def Subtract_Bj(self):
        if self.df_apt is not None:
            self.PeriodicTableCustom.show()
            if self.PeriodicTableCustom.exec() == 1:
                element_dict['charge'] = []
                element_dict['charge'].append(self.PeriodicTableCustom.emit_charge())
                text = ''
                for i in range(len(element_dict['ion'])):
                    text = text + str(element_dict['ion'][i]) + common.subscript(str(element_dict['num'][i]))
                if text:
                    charge = str(element_dict['charge'][0])
                    charge = '(' + charge + ')'
                    text = text + charge
                if self.df_apt[self.df_apt['ION'] == text].shape[0] > 0:
                    prev_text = self.lineEdit_4.text()
                    list_text = prev_text.split(",")
                    if text in list_text:
                        self.df_apt.loc[self.df_apt['ION'] == text, 'status_Bj'] = False
                        text = text + ','
                        prev_text = prev_text.replace(text, '')
                        self.lineEdit_4.setText(prev_text)
                    else:
                        common.show_message("The entered element/ion is not added to the list yet")

                else:
                    common.show_message("The entered element/ion does not exist in the input dataframe")

        else:
            common.show_message("Add the H5 data before addition of Bj/Cl ions")

    # Function to add a specified Cl from the list and show it on the line_edit widget
    def add_Cl(self):
        # text = "Gd₁(2+)"
        # self.df_apt.loc[self.df_apt['ION'] == text, 'status_Cl'] = True
        # self.lineEdit_5.setText(text + ",")

        if self.df_apt is not None:
            self.PeriodicTableCustom.show()
            if self.PeriodicTableCustom.exec() == 1:
                element_dict['charge'] = []
                element_dict['charge'].append(self.PeriodicTableCustom.emit_charge())
                text = ''
                for i in range(len(element_dict['ion'])):
                    text = text + str(element_dict['ion'][i]) + common.subscript(str(element_dict['num'][i]))
                if text:
                    charge = str(element_dict['charge'][0])
                    charge = '(' + charge + ')'
                    text = text + charge
                if self.df_apt[self.df_apt['ION'] == text].shape[0] > 0:
                    prev_text = self.lineEdit_5.text()
                    self.lineEdit_5.setText(prev_text + text + ",")
                    self.df_apt.loc[self.df_apt['ION'] == text, 'status_Cl'] = True

                else:
                    prev_text = self.lineEdit_5.text()
                    self.lineEdit_5.setText(prev_text + "N.A.,")
                    common.show_message("The entered element/ion does not exist in the input dataframe")
        else:
            common.show_message("Add the H5 data before addition of Bj/Cl ions")

    # Function to subtract a specified Cl from the list and show it on the line_edit widget
    def Subtract_Cl(self):
        if self.df_apt is not None:
            self.PeriodicTableCustom.show()
            if self.PeriodicTableCustom.exec() == 1:
                element_dict['charge'] = []
                element_dict['charge'].append(self.PeriodicTableCustom.emit_charge())
                text = ''
                for i in range(len(element_dict['ion'])):
                    text = text + str(element_dict['ion'][i]) + common.subscript(str(element_dict['num'][i]))
                if text:
                    charge = str(element_dict['charge'][0])
                    charge = '(' + charge + ')'
                    text = text + charge
                if self.df_apt[self.df_apt['ION'] == text].shape[0] > 0:
                    prev_text = self.lineEdit_5.text()
                    list_text = prev_text.split(",")
                    if text in list_text:
                        self.df_apt.loc[self.df_apt['ION'] == text, 'status_Cl'] = False
                        text = text + ','
                        prev_text = prev_text.replace(text, '')
                        self.lineEdit_5.setText(prev_text)
                    else:
                        common.show_message("The entered element/ion is not added to the list yet")

                else:
                    common.show_message("The entered element/ion does not exist in the input dataframe")

        else:
            common.show_message("Add the H5 data before addition of Bj/Cl ions")

    def set_default_voxels(self):
        self.df_apt["voxel_x"] = np.nan
        self.df_apt["voxel_y"] = np.nan
        self.df_apt["voxel_z"] = np.nan
        self.df_apt["voxel_a"] = np.nan
        self.df_apt["voxel_b"] = np.nan
        self.df_apt["voxel_c"] = np.nan
        self.df_apt["voxel_number"] = np.nan
        self.df_apt["Solute_per_Voxel"] = np.nan
        self.df_apt['vertex_xyz'] = tuple(self.df_apt[['voxel_x', 'voxel_y', 'voxel_z']].values.tolist())

    # function to voxelize the whole data into eqi-dimensional cubes of input dimension
    def voxelize(self):
        # self.lineEdit_3.setText("3")
        # self.lineEdit_2.setText("0.3")

        if self.df_apt[self.df_apt['status_Bj'] == True].shape[0] > 0:
            pattern_VoxelDia = "^(0|[1-9]\d*)?(\.\d+)?(?<=\d)$"
            rex_pattern_VoxelDia = re.compile(pattern_VoxelDia)
            voxel_cube_dia = self.lineEdit.text()

            if rex_pattern_VoxelDia.match(voxel_cube_dia):
                voxel_cube_dia = float(voxel_cube_dia)
            else:
                common.show_message("a positive float value is expected for Voxel Cube diameter")

            if int(voxel_cube_dia) == 0:
                self.set_default_voxels()

            if voxel_cube_dia > 0 and self.df_apt.shape[0] > 2:
                self.voxel_status = True
                self.df_apt = common.bring_df_to_positive_coord(self.df_apt)
                x_ = np.arange(0., self.df_apt['X'].max(), voxel_cube_dia)
                y_ = np.arange(0., self.df_apt['Y'].max(), voxel_cube_dia)
                z_ = np.arange(0., self.df_apt['Z'].max(), voxel_cube_dia)

                voxel_num = 0
                center = []

                for i in x_:
                    for j in y_:
                        for k in z_:
                            edge = np.array([float(i), j, k])
                            diagonal_edge = np.array([i + voxel_cube_dia, j + voxel_cube_dia, k + voxel_cube_dia], )
                            center_of_voxel = (edge + diagonal_edge) / 2
                            center += [[i, j, k, edge, diagonal_edge, voxel_num, center_of_voxel]]
                            voxel_num += 1

                self.df_apt["voxel_x"] = (self.df_apt["X"] // voxel_cube_dia) * voxel_cube_dia
                self.df_apt["voxel_y"] = (self.df_apt["Y"] // voxel_cube_dia) * voxel_cube_dia
                self.df_apt["voxel_z"] = (self.df_apt["Z"] // voxel_cube_dia) * voxel_cube_dia
                self.df_apt['vertex_xyz'] = tuple(self.df_apt[['voxel_x', 'voxel_y', 'voxel_z']].values.tolist())
                self.df_apt["voxel_number"] = (self.df_apt["X"] // voxel_cube_dia) * (len(z_) * len(y_)) + (
                        self.df_apt["Y"] // voxel_cube_dia) * len(z_) + (self.df_apt["Z"] // voxel_cube_dia)

                list_solute = self.df_apt[self.df_apt['status_Bj'] == True]['ION'].unique()
                num_solute_total = 0
                for solute in list_solute:
                    df_apt_solute = self.df_apt[self.df_apt["ION"] == solute]
                    nodes, inv, counts = np.unique(df_apt_solute["voxel_number"], return_inverse=True,
                                                   return_counts=True)
                    count = counts[inv]
                    df_apt_solute["Solute_per_Voxel"] = count
                    dict_voxelno_soluteno = pd.Series(df_apt_solute.Solute_per_Voxel.values,
                                                      index=df_apt_solute.voxel_number).to_dict()
                    num_solute = self.df_apt["voxel_number"].map(dict_voxelno_soluteno)
                    num_solute = num_solute.fillna(0)
                    num_solute_total = num_solute_total + num_solute

                self.df_apt["Solute_per_Voxel"] = num_solute_total
                self.df_apt = self.df_apt.sort_values(by=['Solute_per_Voxel'], ascending=False)
                self.df_apt['voxel_a'] = voxel_cube_dia
                self.df_apt['voxel_b'] = voxel_cube_dia
                self.df_apt['voxel_c'] = voxel_cube_dia

        else:
            common.show_message("Enter at least one solute Ion in Species_Bj")

    # function to calculate SRO and save the output as hdf file corresponding to each shell
    def Calculate_SRO(self):
        self.gm_sro_array = []
        if self.voxel_status is False:
            self.set_default_voxels()

        def calculate_dist(row, p_iter):
            p2 = (row['X'], row['Y'], row['Z'])
            distance_sq = ((p_iter[0] - p2[0]) ** 2) + ((p_iter[1] - p2[1]) ** 2) + ((p_iter[2] - p2[2]) ** 2)
            return distance_sq

        def shortest_distance(x1, y1, z1, a, b, c, d):
            """ co-ordinate of the given point be (x1, y1, z1) and
            equation of the plane be given by the equation a * x + b * y + c * z + d = 0 """

            d = abs((a * x1 + b * y1 + c * z1 + d))
            e = (math.sqrt(a * a + b * b + c * c))
            return d / e

        # The function checks if the coordinate point in the input row is outside the voxel box after adding r_star
        # it is used to then merge the intersecting voxels based on True values
        def near_edge_solute_check(df_row, r_star):
            x1 = df_row['X'].iloc[0]
            y1 = df_row['X'].iloc[0]
            z1 = df_row['X'].iloc[0]

            voxel_x1 = df_row['voxel_x'].iloc[0]
            voxel_y1 = df_row['voxel_y'].iloc[0]
            voxel_z1 = df_row['voxel_z'].iloc[0]
            voxel_x2 = voxel_x1 + df_row['voxel_a'].iloc[0]
            voxel_y2 = voxel_x1 + df_row['voxel_b'].iloc[0]
            voxel_z2 = voxel_x1 + df_row['voxel_c'].iloc[0]

            # lower YZ plane
            d1 = shortest_distance(x1, y1, z1, a=1, b=0, c=0, d=-(voxel_x1))
            # lower XZ plane
            d2 = shortest_distance(x1, y1, z1, a=0, b=1, c=0, d=-(voxel_y1))
            # lower XY plane
            d3 = shortest_distance(x1, y1, z1, a=0, b=0, c=1, d=-(voxel_z1))
            # greater YZ plane
            d4 = shortest_distance(x1, y1, z1, a=1, b=0, c=0, d=-(voxel_x2))
            # greater XZ plane
            d5 = shortest_distance(x1, y1, z1, a=0, b=1, c=0, d=-(voxel_y2))
            # greater XY plane
            d6 = shortest_distance(x1, y1, z1, a=0, b=0, c=1, d=-(voxel_z2))

            dist_array = [d1, d2, d3, d4, d5, d6]
            min_dist = np.min(dist_array)

            if min_dist <= r_star / 2:
                return True

            return False

        def get_combined_df_voxel(df_row, df_apt_xyz_copy_solute, r_star):
            dist_origin_h = df_row["dist_origin"].values[0] + r_star
            dist_origin_l = df_row["dist_origin"].values[0] - r_star
            C = df_row["X"].values[0] + (3 * r_star)
            D = df_row["X"].values[0] - (3 * r_star)
            E = df_row["Y"].values[0] + (3 * r_star)
            F = df_row["Y"].values[0] - (3 * r_star)
            G = df_row["Z"].values[0] + (3 * r_star)
            H = df_row["Z"].values[0] - (3 * r_star)

            df1 = df_apt_xyz_copy_solute.loc[(df_apt_xyz_copy_solute['dist_origin'] >= dist_origin_l) &
                                             (df_apt_xyz_copy_solute['dist_origin'] <= dist_origin_h)]
            df1 = df1[(df1['Z'].between(H, G)) & (df1['Y'].between(F, E)) & (df1['X'].between(D, C))]

            possible_voxel_list = df1['voxel_number'].unique()
            df_voxel_list = []
            for i in range(len(possible_voxel_list)):
                df_voxel_list.append(
                    df_apt_xyz_copy_solute[df_apt_xyz_copy_solute['voxel_number'] == possible_voxel_list[i]])

            df_voxel_solute = pd.concat(df_voxel_list)
            return df_voxel_solute

        def get_possible_voxel(df_suns, df_voxel_suns, df_voxel_planets, r_star_new):
            suns_dist_origin = df_suns['dist_origin'].iloc[0]
            possible_dist_origin_outer1 = (suns_dist_origin - r_star_new)
            if (suns_dist_origin - r_star_new) < 0:
                possible_dist_origin_outer1 = 0
            possible_dist_origin_outer2 = (suns_dist_origin + r_star_new)

            df_voxel_suns = df_voxel_suns[df_voxel_suns['dist_origin'].between(possible_dist_origin_outer1,
                                                                               possible_dist_origin_outer2)]
            df_voxel_planets = df_voxel_planets[df_voxel_planets['dist_origin'].between(possible_dist_origin_outer1,
                                                                                        possible_dist_origin_outer2)]

            return df_voxel_suns, df_voxel_planets

        shell_no = self.lineEdit_3.text()
        pattern_shell_no = "^[1-9]\d*$"
        rex_pattern_shell_no = re.compile(pattern_shell_no)

        shell_status = False
        if rex_pattern_shell_no.match(shell_no):
            shell_no = int(shell_no)

            if 0 < shell_no < len(self.r_star_array):
                shell_status = True
            else:
                common.show_message("The number of shells exceeds the maximum number of NN radius generated")
        else:
            common.show_message("a positive integer value is expected for number of shells")

        if shell_status:
            for i in range(0, shell_no):
                r_star_new = self.r_star_array[i]
                update = 0
                self.completed = 0
                self.progressBar.setValue(update)

                if (i - 1) < 0:
                    r_star_old = 0
                else:
                    r_star_old = self.r_star_array[i - 1]

                species_BJ = self.df_apt[self.df_apt['status_Bj'] == True]['ION'].unique()
                species_CL = self.df_apt[self.df_apt['status_Cl'] == True]['ION'].unique()

                if len(species_BJ) > 0 and len(species_BJ) > 0:
                    if len(species_BJ) == len(species_BJ):
                        delta_BJ_CL = 1
                    else:
                        delta_BJ_CL = 0

                    X_CL = 0
                    for cl in species_CL:
                        X_CL = X_CL + self.df_apt[self.df_apt["ION"] == cl].shape[0] / self.df_apt.shape[0]

                    print("Shell no: = ", i)
                    print("Delta_Bj_Cl = ", delta_BJ_CL, " and X_Cl = ", X_CL)

                    progress_bar_max = 0
                    for bj in species_BJ:
                        df_apt_bj = self.df_apt[self.df_apt["ION"] == bj]
                        progress_bar_max = progress_bar_max + df_apt_bj.shape[0]
                    progress_bar_max = len(species_CL) * progress_bar_max

                    self.df_apt['P_BJ_CL'] = 0
                    # suns means central atom and planets means the surrounding atoms
                    for bj in species_BJ:
                        print("Calculating clusters with centre, Bj: ", bj, "...")
                        self.df_apt['P_' + str(bj) + '_CL'] = 0
                        self.df_apt['Tot_N_BJ'] = 0
                        self.df_apt["N_" + str(bj)] = self.df_apt[self.df_apt["ION"] == bj].shape[0]
                        self.df_apt['Tot_N_BJ'] = self.df_apt['Tot_N_BJ'] + self.df_apt["N_" + str(bj)]

                        df_apt_suns = self.df_apt[self.df_apt["ION"] == bj]

                        suns_idx_list = df_apt_suns.index.tolist()

                        for cl in species_CL:
                            print("Calculating clusters with solute, Cl: ", cl, "...")
                            for suns_idx in suns_idx_list:
                                self.completed = (self.completed + 1)
                                update = (self.completed / progress_bar_max) * 100
                                self.progressBar.setValue(update)
                                df_suns = df_apt_suns.loc[[suns_idx]]

                                if pd.isna(df_suns['voxel_number'].iloc[0]):
                                    df_voxel_suns = df_suns
                                    df_voxel_planets = self.df_apt
                                    num_planets = df_voxel_planets.shape[0]

                                else:
                                    edge_status = near_edge_solute_check(df_suns, r_star_new)
                                    if edge_status:
                                        df_voxel_suns = get_combined_df_voxel(df_suns, df_apt_suns, r_star_new)
                                        df_voxel_planets = get_combined_df_voxel(df_suns, self.df_apt, r_star_new)

                                    else:
                                        voxel_num = df_suns['voxel_number'].iloc[0]
                                        df_voxel_suns = df_apt_suns[df_apt_suns['voxel_number'] == voxel_num]
                                        df_voxel_planets = self.df_apt[self.df_apt['voxel_number'] == voxel_num]

                                    df_voxel_suns, df_voxel_planets = get_possible_voxel(df_suns,
                                                                                         df_voxel_suns,
                                                                                         df_voxel_planets,
                                                                                         r_star_new)

                                    num_planets = df_voxel_planets.shape[0]

                                if num_planets != 0:
                                    df_cluster = pd.concat([df_voxel_suns, df_voxel_planets], axis=0)
                                    df_cluster = df_cluster[~df_cluster.index.duplicated(keep='first')]
                                    df_cluster.drop(['voxel_x', 'voxel_y', 'voxel_z', 'vertex_xyz', 'dist_origin_sq',
                                                     'Solute_per_Voxel', 'voxel_a', 'voxel_b', 'voxel_c'], axis=1,
                                                    inplace=True)

                                    p_iter = (df_cluster["X"].loc[suns_idx],
                                              df_cluster["Y"].loc[suns_idx],
                                              df_cluster["Z"].loc[suns_idx])

                                    df_cluster['dist_from_idx'] = df_cluster.apply(
                                        lambda row: calculate_dist(row, p_iter),
                                        axis=1)
                                    df_cluster = df_cluster[df_cluster['dist_from_idx'].between(r_star_old,
                                                                                                r_star_new)]

                                    if suns_idx not in df_cluster.index:
                                        df_cluster_row = df_voxel_suns.loc[[suns_idx]]
                                        df_cluster_row['dist_from_idx'] = 0
                                        df_cluster = pd.concat([df_cluster, df_cluster_row], axis=0)

                                    num = df_cluster[df_cluster["ion"] == cl].shape[0]
                                    if cl == bj:
                                        num = df_cluster[df_cluster["ion"] == cl].shape[0] - 1
                                    den = df_cluster.shape[0] - 1

                                    if den != 0:
                                        prob = num / den
                                        self.df_apt.at[suns_idx, 'P_' + str(bj) + '_' + str(cl)] = prob

                            self.df_apt['P_' + str(bj) + '_CL'] = self.df_apt['P_' + str(bj) + '_CL'] + self.df_apt[
                                'P_' + str(bj) + '_' + str(cl)]

                            # end of cl
                            self.df_apt['P_BJ_CL'] = self.df_apt['P_BJ_CL'] + self.df_apt["N_" + str(bj)] * self.df_apt[
                                'P_' + str(bj) + '_CL']
                        # end of bj
                        self.df_apt['P_BJ_CL'] = self.df_apt['P_BJ_CL'] / self.df_apt['Tot_N_BJ']

                    if (delta_BJ_CL - X_CL) != 0:
                        self.df_apt['GM_SRO'] = (-1) ** (1 + delta_BJ_CL) * (
                                (self.df_apt['P_BJ_CL'] - X_CL) / (delta_BJ_CL - X_CL))
                    else:
                        self.df_apt['GM_SRO'] = np.na

                    dir_name = os.path.basename(os.path.splitext(self.hdf_file)[0]) + "_GMSRO"
                    dir_path = os.path.join(os.getcwd(), "GMSRO_Output", dir_name)

                    if os.path.isdir(dir_path) is False:
                        os.makedirs(dir_path)

                    file_path = os.path.join(dir_path, "shell" + str(i) + ".h5")

                    self.df_apt.drop(['voxel_x', 'voxel_y', 'voxel_z', 'vertex_xyz', 'dist_origin_sq',
                                      'Solute_per_Voxel', 'voxel_a', 'voxel_b', 'voxel_c'], axis=1, inplace=False). \
                        to_hdf(file_path, key='df_apt_xyz_copy_' + str(i), mode='w')

                    GMSRO = self.df_apt["GM_SRO"].mean()
                    self.gm_sro_array.append(GMSRO)

        else:
            common.show_message("number of shells must be between 0 and number of radius in the input chart")

    def plot_GMSRO(self):
        if len(self.gm_sro_array) > 0:
            plt.title('Plot of GM_SRO along shells')
            plt.plot(self.gm_sro_array, "ro-")
            plt.xlabel("Shell number (m)")
            plt.ylabel("GM-SRO parameter")
            plt.xticks(np.arange(0, len(self.gm_sro_array), 1.0))
            plt.show()
        else:
            common.show_message("No GM_SRO values were calculated")


# The class inherits a custom UI layout and has facility to analyse composition from H5 file
# Inherits from Ui_CompositionMap
class CompositionMapDialog(Ui_CompositionMap.Ui_Dialog, QDialog):
    def __init__(self, parent=None):
        super(CompositionMapDialog, self).__init__(parent)
        self.setupUi(self)

        self.pushButton.setEnabled(False)
        self.pushButton_3.setEnabled(False)
        self.df_apt_final = None
        self.hdf_file = None
        self.df_apt = None
        self.decompose_el = None
        self.ion_dict = None
        self.scatter3d = None
        self.colorbar = None

        self.widget.fig = Figure()
        self.widget.canvas = FigureCanvas(self.widget.fig)
        self.widget.axes = self.widget.fig.add_subplot(111, projection='3d')

        self.x_plane_start = 0
        self.x_plane_end = 1
        self.y_plane_start = 0
        self.y_plane_end = 1
        self.z_plane_start = 0
        self.z_plane_end = 1

        self.xs = 0
        self.ys = 0
        self.zs = 0
        self.colors = None

        self.buttonBox.accepted.connect(self.accept)
        self.buttonBox.rejected.connect(self.close)

        # Progress bar
        self.completed = 0
        self.progressBar.setValue(self.completed)

        self.pushButton_1.clicked.connect(self.input_file)  # Input H5 file
        self.pushButton_5.clicked.connect(self.add_ion)  # Add an ion to decompose list
        self.pushButton_6.clicked.connect(self.subtract_ion)  # subtract an ion from decompose list
        self.pushButton_7.clicked.connect(self.add_el)  # Add an element to binary plot list
        self.pushButton_8.clicked.connect(self.subtract_el)  # subtract an element to binary plot list
        self.pushButton_3.clicked.connect(self.cluster_analysis)  # calculate the composition based on cluster analysis
        self.pushButton_4.clicked.connect(self.plot_apt)  # plot the APT based on binary elements

    # The function used to read the H5 file containing binned (mapped) apt data
    def input_file(self):
        self.listWidget.clear()
        self.listWidget_2.clear()
        self.listWidget_3.clear()
        self.listWidget_4.clear()
        # file = ['C:/Users/arjun/Downloads/APT_Code/APT_Project/Output_Analyzed_H5/Cooxide_binned.h5']
        file = QFileDialog.getOpenFileName(self, 'Select HDF File"',
                                           os.getcwd(), "HDF files (*.h5)")

        try:
            self.hdf_file = file[0]
            self.df_apt = pd.read_hdf(self.hdf_file)
            self.df_apt = common.bring_df_to_positive_coord(self.df_apt)
            self.df_apt['status_decompose'] = None
            # each time the previous cluster analysis is replaced, better book-keeping can be optionally added for reuse
            if 'cluster_centre' in self.df_apt.columns:
                del self.df_apt['cluster_centre']
            self.df_apt['cluster_id'] = 0
            self.lcdNumber.setDigitCount(10)
            self.lcdNumber.display(0)
            self.lineEdit_3.setText('3')
            self.lineEdit_2.setText('100')
            self.lineEdit_4.setText('3')
            self.lineEdit_5.setText('0.1')

            self.ion_dict = {}
            for ion in self.df_apt['ION'].unique():
                ion_num = self.df_apt[self.df_apt['ION'] == ion].shape[0]
                self.ion_dict[ion] = ion_num

            for ion in self.df_apt['ION'].unique():
                self.listWidget.addItem(str(ion))

            self.pushButton.setEnabled(True)
            self.pushButton_3.setEnabled(True)
        except:
            self.pushButton.setEnabled(False)
            self.pushButton_3.setEnabled(False)

    # Function to add ion to the decompose list
    def add_ion(self):
        if self.listWidget.currentItem():
            selected_text = self.listWidget.currentItem().text()
            self.df_apt.loc[self.df_apt['ION'] == selected_text, 'status_decompose'] = True
            self.listWidget_2.addItem(selected_text)
            self.listWidget.takeItem(self.listWidget.currentRow())
            ion_num = self.ion_dict[selected_text]
            self.lcdNumber.display(self.lcdNumber.value() + ion_num)

    # Function to subtract ion from the decompose list
    def subtract_ion(self):
        if self.listWidget_2.currentItem():
            selected_text = self.listWidget_2.currentItem().text()
            self.df_apt.loc[self.df_apt['ION'] == selected_text, 'status_decompose'] = False
            self.listWidget.addItem(selected_text)
            self.listWidget_2.takeItem(self.listWidget_2.currentRow())
            ion_num = self.ion_dict[selected_text]
            self.lcdNumber.display(self.lcdNumber.value() - ion_num)

    # Function to add el to the element list
    def add_el(self):
        if self.listWidget_3.currentItem():
            if self.listWidget_4.count() < 2:
                selected_text = self.listWidget_3.currentItem().text()
                self.listWidget_4.addItem(selected_text)
                self.listWidget_3.takeItem(self.listWidget_3.currentRow())

    # Function to subtract el from the element list
    def subtract_el(self):
        if self.listWidget_4.currentItem():
            selected_text = self.listWidget_4.currentItem().text()
            self.listWidget_3.addItem(selected_text)
            self.listWidget_4.takeItem(self.listWidget_4.currentRow())

    # Function to get all the parameter values for clustering after regular expression check from the UI
    def input_parameters(self):
        pattern_NN = "^[0-9]*[1-9][0-9]*$"
        pattern_num_origins = "^[0-9]*[1-9][0-9]*$"
        pattern_leaf = "^[0-9]*[1-9][0-9]*$"
        pattern_critical_radius = "(0|[1-9]\d*)?(\.\d+)?(?<=\d)$"

        rex_pattern_NN = re.compile(pattern_NN)
        rex_pattern_num_origins = re.compile(pattern_num_origins)
        rex_pattern_leaf = re.compile(pattern_leaf)
        rex_pattern_critical_radius = re.compile(pattern_critical_radius)

        num_origins = self.lineEdit_2.text()
        NN = self.lineEdit_3.text()
        leaf = self.lineEdit_4.text()
        critical_radius = self.lineEdit_4.text()

        if rex_pattern_NN.match(NN):
            NN = int(NN)
        else:
            common.show_message("a positive integer value is expected for No: of Nearest Neighbors")

        if rex_pattern_num_origins.match(num_origins):
            num_origins = int(num_origins)
        else:
            common.show_message("a positive integer value is expected for No: of origins")

        if rex_pattern_leaf.match(leaf):
            leaf = int(leaf)
        else:
            common.show_message("a positive integer value is expected for No: of leaves")

        if rex_pattern_critical_radius.match(critical_radius):
            critical_radius = float(critical_radius)
        else:
            common.show_message("a positive float value is expected for critical radius")

        return NN, num_origins, leaf, critical_radius

    # Function to analyse clusters and get composition data based on ions to decompose
    def cluster_analysis(self):
        self.listWidget_3.clear()
        self.listWidget_4.clear()

        def calculate_distance(p1, p2):
            return (p1[0] - p2[0]) ** 2 + (p1[1] - p2[1]) ** 2 + (p1[2] - p2[2]) ** 2

        if self.df_apt[self.df_apt['status_decompose'] == True].shape[0] > 0:
            NN, num_origins, leaf, critical_radius = self.input_parameters()

            self.decompose_el = [str(self.listWidget_2.item(i).text()) for i in range(self.listWidget_2.count())]
            self.df_apt['status_decompose'] = self.df_apt['ION'].apply(
                lambda x: any([i in x for i in self.decompose_el]))
            df_decompose = self.df_apt[self.df_apt['status_decompose'] == True]

            if num_origins > self.lcdNumber.value():
                num_origins = self.lcdNumber.value()

            dict_apt_neighbour = defaultdict(list)
            XYZ_list_complete = np.asarray(
                list(zip(df_decompose['X'].tolist(), df_decompose['Y'].tolist(), df_decompose['Z'].tolist())))
            random_df_apt = df_decompose.iloc[
                np.random.choice(np.arange(df_decompose.shape[0]), size=num_origins, replace=False)]
            XYZ_list_origins = np.asarray(list(zip(random_df_apt.index.values.tolist(), random_df_apt['X'].tolist(),
                                                   random_df_apt['Y'].tolist(), random_df_apt['Z'].tolist())))
            tree = KDTree(XYZ_list_complete[:, 0:3], leaf_size=leaf)

            # If we wish to add voxel splits, then do KD tree for all voxels and find idx based on corresponding tree
            # Merge the voxels if the origin lies near the edge and do kd tree again for such voxels
            # Now during comparison with previous clusters (using df_temp) only compare inside its voxels using voxel id

            i = 0
            self.completed = i
            self.progressBar.setValue(self.completed)
            centre_current_cluster = [0, 0, 0]
            while i < num_origins:
                cluster_id_current = i + 1
                self.completed = ((i + 1) / num_origins) * 100
                self.progressBar.setValue(self.completed)

                origin = [[t for t in XYZ_list_origins[i]][1:4]]
                _, idx = tree.query(origin, k=NN)
                list_neighbours = []
                for j in range(len(idx[0])):
                    list_neighbours.append(idx[0][j])

                for k in range(len(list_neighbours)):
                    flag_if_condition = False
                    flag_else_condition = False
                    flag_freedom_condition = True

                    index = list_neighbours[k]
                    X, Y, Z = XYZ_list_complete[index][0], XYZ_list_complete[index][1], XYZ_list_complete[index][2]

                    if k == 0:
                        centre_current_cluster = [X, Y, Z]
                        current_neighbour = [X, Y, Z]
                        current_index = index

                    else:
                        current_neighbour = [X, Y, Z]
                        current_index = index

                    if index in dict_apt_neighbour['index']:
                        temp_df = pd.DataFrame.from_dict(dict_apt_neighbour, orient='columns')
                        cluster_id_prev = temp_df.loc[temp_df['index'] == index, 'cluster_id'].values[0]
                        temp_df_true_centres = temp_df[temp_df['cluster_centre'] == True]
                        temp_df_true_centres_id = temp_df_true_centres[
                            temp_df_true_centres['cluster_id'] == cluster_id_prev]
                        centre_prev_cluster = [temp_df_true_centres_id['X'].values[0],
                                               temp_df_true_centres_id['Y'].values[0],
                                               temp_df_true_centres_id['Z'].values[0]]

                        curr_distance = calculate_distance(centre_current_cluster, current_neighbour)
                        prev_distance = calculate_distance(centre_prev_cluster, current_neighbour)

                        if curr_distance > prev_distance:
                            flag_if_condition = True

                        else:
                            if curr_distance == 0:
                                # This part has not been updated in jupyter notebook
                                if prev_distance < critical_radius:
                                    break

                            if flag_freedom_condition:
                                flag_else_condition = True

                                if k == 0:
                                    temp_df['cluster_centre'].loc[temp_df['index'] == current_index] = True
                                else:
                                    temp_df['cluster_centre'].loc[temp_df['index'] == current_index] = False

                                temp_df['cluster_id'].loc[temp_df['index'] == current_index] = cluster_id_current

                                dict_apt_neighbour = {'index': temp_df['index'].tolist(), 'X': temp_df['X'].tolist(),
                                                      'Y': temp_df['Y'].tolist(), 'Z': temp_df['Z'].tolist(),
                                                      'cluster_centre': temp_df['cluster_centre'].tolist(),
                                                      'cluster_id': temp_df['cluster_id'].tolist()}

                    if flag_else_condition is False:
                        if flag_if_condition is True:
                            pass

                        else:
                            dict_apt_neighbour['index'].append(index)
                            dict_apt_neighbour['X'].append(X)
                            dict_apt_neighbour['Y'].append(Y)
                            dict_apt_neighbour['Z'].append(Z)

                            if k == 0:
                                dict_apt_neighbour['cluster_centre'].append(True)
                            else:
                                dict_apt_neighbour['cluster_centre'].append(False)

                            dict_apt_neighbour['cluster_id'].append(cluster_id_current)

                i = i + 1

            df_apt_neighbour = pd.DataFrame.from_dict(dict_apt_neighbour, orient='columns')

            merged = pd.merge(self.df_apt, df_apt_neighbour, how='left', left_index=True, right_on='index')
            merged['cluster_id'] = np.max(merged[['cluster_id_x', 'cluster_id_y']], axis=1)
            merged['X'] = np.max(merged[['X_x', 'X_y']], axis=1)
            merged['Y'] = np.max(merged[['Y_x', 'Y_y']], axis=1)
            merged['Z'] = np.max(merged[['Z_x', 'Z_y']], axis=1)
            merged = merged.drop(labels=['index', 'cluster_id_x', 'cluster_id_y', 'X_x', 'X_y',
                                         'Y_x', 'Y_y', 'Z_x', 'Z_y'], axis=1)
            self.df_apt_final = merged.sort_values(by=['cluster_id'])

            df_decompose = self.df_apt_final[self.df_apt_final['status_decompose'] == True]
            df_cluster = df_decompose[df_decompose.cluster_centre.notna()]
            list_el = np.unique([*itertools.chain.from_iterable(df_cluster.ion)])
            for el in list_el:
                self.listWidget_3.addItem(str(el))
        else:
            show_message("Pick at least one ion to decompose from the list before cluster analysis")

    # The plot function to plot either of the elements or their gradient
    def plot_apt(self):
        if self.listWidget_4.count() == 2:
            df_decompose = self.df_apt_final[self.df_apt_final['status_decompose'] == True]
            df_cluster = df_decompose[df_decompose.cluster_centre.notna()]
            df_cluster = df_cluster.drop(labels=['MN_Ratio', 'peak_no', 'peak_MNRatio',
                                                 'peak_max_cutoff_width', 'XYZ_total_count', 'status_decompose'],
                                         axis=1)
            self.decompose_el = [str(self.listWidget_4.item(i).text()) for i in range(self.listWidget_4.count())]

            def return_val(row, el):
                ion = list(row['ion'])
                num = list(row['num'])
                if el in ion:
                    idx = ion.index(el)
                    return float(num[idx])
                else:
                    return 0

            def return_conc(row):
                if row['element1_conc'] == 0 and row['element2_conc'] == 0:
                    conc = np.na
                    label = np.na

                elif row['element1_conc'] == 0:
                    conc = row['element2_conc']
                    label = self.decompose_el[1]

                elif row['element2_conc'] == 0:
                    conc = row['element1_conc']
                    label = self.decompose_el[0]

                else:
                    conc = row['element1_conc'] / (row['element1_conc'] + row['element2_conc'])
                    label = 'grad'

                return pd.Series([conc, label])

            df_cluster["element1_val"] = df_cluster.apply(lambda row: return_val(row, self.decompose_el[0]), axis=1)
            df_cluster["element2_val"] = df_cluster.apply(lambda row: return_val(row, self.decompose_el[1]), axis=1)
            df_temp1 = df_cluster.groupby(['cluster_id'])["element1_val"].sum().reset_index(name="element1_conc")
            df_temp2 = df_cluster.groupby(['cluster_id'])["element2_val"].sum().reset_index(name="element2_conc")
            df_cluster = pd.merge(df_cluster, df_temp1, how='left', left_on=['cluster_id'], right_on=['cluster_id'])
            df_cluster = pd.merge(df_cluster, df_temp2, how='left', left_on=['cluster_id'], right_on=['cluster_id'])
            df_cluster[['conc_grad', 'label']] = df_cluster.apply(lambda row: return_conc(row), axis=1)

            plot_label = None
            if self.radioButton.isChecked():
                plot_label = self.decompose_el[0]

            elif self.radioButton_2.isChecked():
                plot_label = self.decompose_el[1]

            elif self.radioButton_3.isChecked():
                plot_label = 'grad'

            df_cluster_plot = df_cluster[df_cluster['label'] == plot_label][['X', 'Y', 'Z', 'conc_grad']]

            if df_cluster_plot.shape[0] > 2:
                layout = QVBoxLayout(self.widget)
                layout.addWidget(self.widget.canvas)
                self.x_plane_start = df_cluster_plot["X"].min()
                self.x_plane_end = df_cluster_plot["X"].max()
                self.y_plane_start = df_cluster_plot["Y"].min()
                self.y_plane_end = df_cluster_plot["Y"].max()
                self.z_plane_start = df_cluster_plot["Z"].min()
                self.z_plane_end = df_cluster_plot["Z"].max()

                self.xs = df_cluster_plot['X'].astype(float)
                self.ys = df_cluster_plot['Y'].astype(float)
                self.zs = df_cluster_plot['Z'].astype(float)
                self.colors = df_cluster_plot['conc_grad'].astype(float)

                if self.scatter3d:
                    self.scatter3d.remove()
                del self.scatter3d
                self.scatter3d = None

                cm = plt.cm.get_cmap('RdYlBu')
                self.scatter3d = self.widget.axes.scatter(self.xs, self.ys, self.zs, c=self.colors,
                                                          cmap=cm, vmin=0, vmax=self.colors.max(),
                                                          s=30, label=plot_label, marker='.')

                self.widget.fig.subplots_adjust(0.2, 0.2, 0.8, 0.8)  # left,bottom,right,top
                if self.colorbar is None:
                    self.colorbar = plt.colorbar(self.scatter3d, ax=self.widget.axes)
                else:
                    self.colorbar.mappable.set_clim(0, self.colors.max())

                self.widget.axes.set_xlim3d(self.x_plane_start, self.x_plane_end)
                self.widget.axes.set_ylim3d(self.y_plane_start, self.y_plane_end)
                self.widget.axes.set_zlim3d(self.z_plane_start, self.z_plane_end)
                self.widget.axes.legend()
                self.widget.axes.set_xlabel('X Axis')
                self.widget.axes.set_ylabel('Y Axis')
                self.widget.axes.set_zlabel('Z Axis')
                self.widget.fig.canvas.draw()
                self.widget.fig.tight_layout()

            else:
                show_message("Not sufficient clusters detected to plot")

        else:
            show_message("Two elements are needed for plotting")


class MainWindow(Ui_APTMainWindow.Ui_MainWindow, QMainWindow):
    def __init__(self, parent=None):
        super(MainWindow, self).__init__()
        self.setupUi(self)
        self.setWindowTitle("APT Analyzer - Version 1.0")

        # Self Variables for included Windows
        self.plot_window = None
        self.input_table = None
        self.mono_layer = None
        self.abstract_layer = None
        self.SRO = None
        self.composition_map = None

        # Self Variables for dataframe
        self.model = None
        self.df_apt = None
        self.df_el = None
        self.df_apt_final = None

        # Self Variables for functions
        self.mat_file = None
        self.completed = 0

        # Progress bar
        self.progressBar.setValue(self.completed)

        # Context Menu and analysis Tabs
        self.actionOpenMatFile.triggered.connect(self.input_file)
        self.actionMono_Layer_Analysis.triggered.connect(self.Mono_Layer_Analysis)
        self.actionAbstract_Layer_Analysis.triggered.connect(self.Abstract_Layer_Analysis)
        self.actionGM_SRO.triggered.connect(self.GM_SRO)
        self.actionComposition_Mapping.triggered.connect(self.Composition_Mapping)
        self.actionExit.triggered.connect(sys.exit)
        self.start_button_status()
        self.button_operations()

    def button_operations(self):
        """
        btn_view_df = Prints the current dataframe
        btn_plot_hist = Plots Histogram
        btn_input_elem = Input the elements table
        btn_start_bin = Peak mapping using input table and bin size
        btn_export_hdf = Save the final df_apt as HDF file for further analysis
        btn_exit = Exit
        btn_view_peak_df = View the df elements here
        btn_view_final_df = View the whole df apt after binning here
        """
        self.btn_plot_hist.clicked.connect(self.plot_hist)
        self.btn_input_elem.clicked.connect(self.input_elements)
        self.btn_view_df.clicked.connect(self.view_df_apt)
        self.btn_start_bin.clicked.connect(self.start_binning)
        self.btn_view_peak_df.clicked.connect(self.view_df_el)
        self.btn_view_final_df.clicked.connect(self.view_df_apt_final)
        self.btn_export_hdf.clicked.connect(self.export_hdf)
        self.btn_exit.clicked.connect(QtCore.QCoreApplication.instance().quit)

    def start_button_status(self):
        """   set status of the buttons when program starts before input_file """

        self.tableView.horizontalHeader().setStretchLastSection(True)
        self.btn_view_df.setEnabled(False)
        self.btn_plot_hist.setEnabled(False)
        self.btn_start_bin.setEnabled(False)
        self.btn_view_peak_df.setEnabled(False)
        self.btn_view_final_df.setEnabled(False)
        self.btn_export_hdf.setEnabled(False)

    def input_file(self):
        """ Shows the Folder browser dialog to input the .mat file  """

        file = QFileDialog.getOpenFileName(self, 'Select Matlab File"',
                                           os.getcwd(), "Mat files (*.mat)")
        try:
            self.mat_file = file[0]
            self.df_apt = common.read_data(self.mat_file)
            self.start_button_status()
            self.btn_view_df.setEnabled(True)
            self.btn_plot_hist.setEnabled(True)
            self.btn_start_bin.setEnabled(True)

        except:
            self.btn_view_df.setEnabled(False)
            self.btn_start_bin.setEnabled(False)

    # The function does cluster analysis through the dataset to color map required elements
    def Composition_Mapping(self):
        self.composition_map = CompositionMapDialog()
        self.composition_map.show()

    # The function that calls Mono layer window to do the analysis part and get ion counts
    def Mono_Layer_Analysis(self):
        self.mono_layer = MonoLayerDialog()
        self.mono_layer.show()

    # The function that calls Abstract layer window to do the analysis part and get ion counts
    def Abstract_Layer_Analysis(self):
        self.abstract_layer = AbstractLayerDialog()
        self.abstract_layer.show()

    # Opens the SRO Dialog
    def GM_SRO(self):
        self.SRO = SRODialog()
        self.SRO.show()

    # To export the final binned and mapped apt dataset as hdf file for further analysis
    def export_hdf(self):
        file = QFileDialog.getSaveFileName(self, 'Select/Create HDF File"',
                                           os.getcwd(), "HDF files (*.h5)")
        if len(file[0]) > 1:
            float_columns = ['X', 'Y', 'Z', 'MN_Ratio', 'peak_MNRatio', 'peak_max_cutoff_width']
            int_columns = ['peak_no', 'XYZ_total_count']

            self.df_apt_final.loc[:, float_columns] = self.df_apt_final[float_columns].applymap(float)
            self.df_apt_final.loc[:, int_columns] = self.df_apt_final[int_columns].applymap(int)

            self.df_apt_final.to_hdf(file[0], key='df_apt_final', mode='w')

    # To view the columns of the original file input
    def view_df_apt(self):
        self.df_apt = common.read_data(self.mat_file)
        # self.model = DataFrameModel(self.df_apt.tail())

        self.model = QStandardItemModel()
        table_title = list(self.df_apt.columns.values)
        table_title.insert(0, "row_idx")
        self.model.setHorizontalHeaderLabels(table_title)

        index, row_data = self.df_apt.tail().index.values, self.df_apt.tail().to_numpy()
        data = np.hstack((index.reshape(len(index), 1), row_data))

        for line in data:
            row = []
            for name in line:
                item = QStandardItem(str(name))
                item.setEditable(False)
                row.append(item)
            self.model.appendRow(row)

        self.tableView.setModel(self.model)
        for i in range(self.df_apt.shape[1]):
            self.tableView.horizontalHeader().setSectionResizeMode(i, QHeaderView.ResizeToContents)

    # To view the columns of the final binned and mapped apt dataset
    def view_df_apt_final(self):
        # self.model = DataFrameModel(self.df_apt_final.head())

        self.model = QStandardItemModel()
        table_title = list(self.df_apt_final.columns.values)
        table_title.insert(0, "row_idx")
        self.model.setHorizontalHeaderLabels(table_title)
        index, row_data = self.df_apt_final.tail().index.values, self.df_apt_final.tail().to_numpy()
        data = np.hstack((index.reshape(len(index), 1), row_data))

        for line in data:
            row = []
            for name in line:
                item = QStandardItem(str(name))
                item.setEditable(False)
                row.append(item)
            self.model.appendRow(row)

        self.tableView.setModel(self.model)
        for i in range(self.df_apt.shape[1]):
            self.tableView.horizontalHeader().setSectionResizeMode(i, QHeaderView.ResizeToContents)

    # To view the final columns of all elements that were provided as input from the table
    def view_df_el(self):
        self.model = DataFrameModel(self.df_el)
        self.tableView.setModel(self.model)
        for i in range(self.df_el.shape[1]):
            self.tableView.horizontalHeader().setSectionResizeMode(i, QHeaderView.ResizeToContents)

    # To plot the histogram data given the start and end MN ratio, helps in checking for peaks
    def plot_hist(self):
        if self.lineEdit.text() == "":
            xlim_left = None
        else:
            string_xlim_left = self.lineEdit.text()
            xlim_left = re.sub("\D", "", string_xlim_left)
            xlim_left = float(string_xlim_left)

        if self.lineEdit_2.text() == "":
            xlim_right = None
        else:
            string_xlim_right = self.lineEdit_2.text()
            xlim_right = re.sub("\D", "", string_xlim_right)
            xlim_right = float(string_xlim_right)

        if self.lineEdit_3.text() == "":
            bins = None
        else:
            string_bins = self.lineEdit_3.text()
            bins = re.sub("\D", "", string_bins)
            bins = float(string_bins)

        if xlim_left is None or xlim_right is None or bins is None:
            show_message("No Parameters for Histogram")

        elif bins > 1:
            show_message("Enter a valid bin size (0 - 1)")

        else:
            num_bins = int((1 / bins) * np.max(self.df_apt["MN_Ratio"]))
            _freq, _bins = np.histogram(self.df_apt["MN_Ratio"], bins=num_bins)
            df_apt_hist = pd.DataFrame(list(zip(_bins[:-1], _bins[:-1], _freq)),
                                       columns=['bin_lower', 'bin_upper', 'freq'])

            subset_df_apt_hist = df_apt_hist[df_apt_hist["bin_lower"] > float(xlim_left)]
            subset_df_apt_hist = subset_df_apt_hist[subset_df_apt_hist["bin_lower"] < float(xlim_right)]

            self.plot_window = HistogramWindow(subset_df_apt_hist)
            self.plot_window.show()

    # Open the input element table to input the element information and cutoff values
    def input_elements(self):
        self.input_table = InputElementTable()
        self.input_table.show()

    # The binning operation which maps the peak according to the input table
    def start_binning(self):
        list_of_dict = []
        self.completed = 0
        self.progressBar.setValue(self.completed)

        for key1 in element_dict_array:
            for key2 in cutoff_dict_array:
                if key2 == key1:
                    x = element_dict_array[key1]
                    y = cutoff_dict_array[key2]
                    z = {**x, **y}
                    list_of_dict.append(z)

        self.df_el = pd.DataFrame(list_of_dict)
        if self.df_el.empty:
            common.show_message("Enter the Elements Table..")
        if not self.df_el.empty:
            none_values = self.df_el.isnull().values.sum()
            if none_values > 0:
                common.show_message("Input elements and complete the table to start binning..")
            else:
                cutoff_bins = [float(i) for i in self.df_el['cutoff_bin'].values]
                peak_MNRatio = [float(i) for i in self.df_el['peak_MNRatio'].values]
                peak_width = [float(i) for i in self.df_el['peak_width'].values]
                cutoff_height = [int(i) for i in self.df_el['cutoff_height'].values]
                cutoff_width = [int(i) for i in self.df_el['cutoff_width'].values]

                class_iter = 0
                peak_no_max = 0
                dict_sub_df = {}
                dict_sub_df_hist = {}
                dict_sub_df_merged = {}

                while class_iter < self.df_el.shape[0]:
                    self.completed = (class_iter / self.df_el.shape[0]) * 100.0
                    self.progressBar.setValue(self.completed)
                    reciprocal_bins = 1 / cutoff_bins[class_iter]
                    MNRatio_start = float(peak_MNRatio[class_iter] - peak_width[class_iter] / 2.0)
                    MNRatio_end = float(peak_MNRatio[class_iter] + peak_width[class_iter] / 2.0)
                    dict_sub_df[class_iter] = self.df_apt[self.df_apt["MN_Ratio"].between(MNRatio_start, MNRatio_end)]
                    num_bins = int(reciprocal_bins * (MNRatio_end - MNRatio_start))
                    bin_intervels = np.linspace(MNRatio_start, MNRatio_end, num_bins)
                    freq, bins = np.histogram(dict_sub_df[class_iter]["MN_Ratio"], bins=bin_intervels)

                    def truncate(f, n):
                        return math.floor(f * 10 ** n) / 10 ** n

                    for ib in range(len(bins)):
                        bins[ib] = truncate(bins[ib], np.log10(reciprocal_bins))

                    # freq will become peaks if it satisfy the given height and width conditions
                    dict_sub_df_hist[class_iter] = pd.DataFrame(list(zip(bins[:-1],
                                                                         bins[:-1] +
                                                                         cutoff_bins[class_iter],
                                                                         freq)),
                                                                columns=['bin_lower', 'bin_upper', 'freq'])

                    dict_sub_df_hist[class_iter]['Range_Cutoff_H_Value'] = cutoff_height[class_iter]
                    dict_sub_df_hist[class_iter]['Range_Cutoff_W_Value'] = cutoff_width[class_iter]

                    def check_cutoff_height(row):
                        if row['freq'] > row['Range_Cutoff_H_Value']:
                            return True
                        return False

                    dict_sub_df_hist[class_iter]['cutoff_height_status'] = dict_sub_df_hist[class_iter].apply(
                        lambda row: check_cutoff_height(row), axis=1)
                    dict_sub_df_hist[class_iter]['subgroup'] = \
                        (dict_sub_df_hist[class_iter]['cutoff_height_status'] !=
                         dict_sub_df_hist[class_iter]['cutoff_height_status'].shift(1)).cumsum()
                    dict_sub_df_hist[class_iter]['subgroup_freq'] = dict_sub_df_hist[class_iter].groupby('subgroup')[
                        'subgroup'].transform('count')

                    def check_cutoff_width(row):
                        if row['subgroup_freq'] > row['Range_Cutoff_W_Value'] and row['cutoff_height_status'] is True:
                            return True
                        return False

                    dict_sub_df_hist[class_iter]['cutoff_height_width_status'] = dict_sub_df_hist[class_iter].apply(
                        lambda row: check_cutoff_width(row), axis=1)
                    peak_cutoff_status = dict_sub_df_hist[class_iter]['cutoff_height_width_status'].tolist()

                    if class_iter == 0:
                        peak_no_iter = 0
                        peak_no_max = 0
                        peak_cutoff_status[0] = False
                    else:
                        peak_cutoff_status[0] = False
                        for class_idx in range(class_iter):
                            peak_no_max_idx = dict_sub_df_hist[class_idx]['peak_no'].max()
                            peak_no_max = max(peak_no_max, peak_no_max_idx)

                        peak_no_iter = peak_no_max

                    peak = np.zeros(len(peak_cutoff_status), dtype=np.int64)
                    for j in range(len(peak_cutoff_status)):
                        if peak_cutoff_status[j] is False:
                            peak[j] = 0
                        if peak_cutoff_status[j] is True:
                            if j > 0 and peak_cutoff_status[j - 1] is False:
                                peak_no_iter = peak_no_iter + 1
                            peak[j] = int(peak_no_iter)

                    dict_sub_df_hist[class_iter]['peak_no'] = peak
                    MN_Ratio = dict_sub_df[class_iter].MN_Ratio.values
                    bin_lower = dict_sub_df_hist[class_iter].bin_lower.values
                    bin_upper = dict_sub_df_hist[class_iter].bin_upper.values
                    ii, jj = np.where((MN_Ratio[:, None] >= bin_lower) & (MN_Ratio[:, None] <= bin_upper))
                    dict_sub_df_merged[class_iter] = pd.DataFrame(
                        np.column_stack([dict_sub_df[class_iter].values[ii], dict_sub_df_hist[class_iter].values[jj]]),
                        columns=dict_sub_df[class_iter].columns.append(dict_sub_df_hist[class_iter].columns))

                    class_iter = class_iter + 1

                df_apt_merged = pd.concat(dict_sub_df_merged.values(), ignore_index=True)

                max_peak_no = df_apt_merged['peak_no'].max() + 1
                dict_peak_min_max_count = dict()

                if max_peak_no > 1:
                    for peak_id in range(1, max_peak_no):
                        temp_df_apt = df_apt_merged[df_apt_merged["peak_no"] == peak_id]
                        max_cutoff_width = temp_df_apt.subgroup_freq.mode()[0]
                        peak_total_count = temp_df_apt.shape[0]
                        dict_peak_min_max_count[peak_id] = peak_id, temp_df_apt['MN_Ratio'].min(), \
                                                           temp_df_apt['MN_Ratio'].max(), max_cutoff_width, \
                                                           peak_total_count

                    def check_val_in_dict(peak_input, dict_peak_min_max_count):
                        non_peaks = 0

                        for key in dict_peak_min_max_count.keys():
                            if dict_peak_min_max_count[key][1] < peak_input < dict_peak_min_max_count[key][2]:
                                return pd.Series([int(dict_peak_min_max_count[key][0]),
                                                  dict_peak_min_max_count[key][3], dict_peak_min_max_count[key][4]])
                        message = "No peaks were observed with input cutoff_conditions at MN_Ratio: " + str(peak_input)

                        def scarce_element(peak_input):
                            tolerance = 0
                            cutoff_bin = None

                            for key in cutoff_dict_array:
                                if cutoff_dict_array[key]['peak_MNRatio'] is not None:
                                    if float(cutoff_dict_array[key]['peak_MNRatio']) == float(peak_input):
                                        tolerance = float(cutoff_dict_array[key]['peak_width'])
                                        cutoff_bin = float(cutoff_dict_array[key]['cutoff_bin'])
                                        break

                            MNRatio_start = float(peak_input - tolerance / 2.0)
                            MNRatio_end = float(peak_input + tolerance / 2.0)
                            df_apt_scarce = self.df_apt[self.df_apt["MN_Ratio"].between(MNRatio_start, MNRatio_end)]
                            reciprocal_bins = 1 / cutoff_bin
                            num_bins = int(reciprocal_bins * (MNRatio_end - MNRatio_start))
                            bin_intervels = np.linspace(MNRatio_start, MNRatio_end, num_bins)
                            freq, bins = np.histogram(df_apt_scarce["MN_Ratio"], bins=bin_intervels)

                            hist_apt_scarce = pd.DataFrame(list(zip(bins[:-1], bins[:-1] + cutoff_bin, freq)),
                                                           columns=['bin_lower', 'bin_upper', 'freq'])

                            hist_apt_scarce['bin_lower'] = hist_apt_scarce['bin_lower'].round(
                                int(np.log10(reciprocal_bins)))
                            hist_apt_scarce['bin_upper'] = hist_apt_scarce['bin_upper'].round(
                                int(np.log10(reciprocal_bins)))

                            self.model = DataFrameModel(hist_apt_scarce)
                            self.tableView.setModel(self.model)
                            for i in range(hist_apt_scarce.shape[1]):
                                self.tableView.horizontalHeader().setSectionResizeMode(i, QHeaderView.ResizeToContents)

                        common.show_message(message, btn1=True, btn1_name="View Scarce Element Table",
                                            btn1_fun=lambda: scarce_element(peak_input), btn2=True, btn2_name="OK",
                                            btn2_fun=lambda: None)

                        return pd.Series([float('nan'), float('nan'), 0])

                    self.df_el[['peak_id', 'peak_max_cutoff_width', 'XYZ_total_count']] = self.df_el[
                        'peak_MNRatio'].apply(lambda x: check_val_in_dict(float(x), dict_peak_min_max_count))

                    self.df_apt_final = pd.merge(df_apt_merged, self.df_el, how='left', left_on='peak_no',
                                                 right_on='peak_id')
                    self.df_apt_final = self.df_apt_final.dropna(subset=['ion'])
                    self.df_apt_final = self.df_apt_final.drop(
                        columns=['bin_lower', 'bin_upper', 'freq', 'Range_Cutoff_H_Value',
                                 'Range_Cutoff_W_Value', 'subgroup', 'subgroup_freq',
                                 'cutoff_height_status', 'cutoff_height_width_status',
                                 'peak_width', 'cutoff_bin', 'cutoff_height',
                                 'cutoff_width', 'peak_id'])

                df_peak_min_max_count = pd.DataFrame(columns=['peak_id', 'min_MN', 'max_MN', 'peak_max_cutoff_width'])
                for key in dict_peak_min_max_count.keys():
                    df_peak_min_max_count = df_peak_min_max_count.append(
                        pd.Series({'peak_id': dict_peak_min_max_count[key][0],
                                   'min_MN': dict_peak_min_max_count[key][1],
                                   'max_MN': dict_peak_min_max_count[key][2],
                                   'peak_max_cutoff_width': dict_peak_min_max_count[key][3]}),
                        ignore_index=True)

                self.df_apt_final = self.df_apt_final.sort_values(by=['MN_Ratio'], ignore_index=True)

                def change_ion_name(row):
                    ion_array = row['ion']
                    text = ''
                    for ii in range(len(ion_array)):
                        text = text + ion_array[ii] + common.subscript(row['num'][ii])
                        if ii == len(ion_array) - 1:
                            text = text + '(' + row['charge'][0] + ')'
                    return text

                self.df_el['ION'] = self.df_el.apply(lambda row: change_ion_name(row), axis=1)
                self.df_apt_final['ION'] = self.df_apt_final.apply(lambda row: change_ion_name(row), axis=1)

                self.df_el = self.df_el.drop(columns=['ion', 'num', 'charge', 'peak_width',
                                                      'cutoff_bin', 'cutoff_height', 'cutoff_width'])
                cols = ['ION', 'mass', 'peak_MNRatio', 'peak_id', 'peak_max_cutoff_width', 'XYZ_total_count']
                self.df_el = self.df_el[cols]

                self.completed = 100
                self.progressBar.setValue(self.completed)

                self.btn_view_peak_df.setEnabled(True)
                self.btn_view_final_df.setEnabled(True)
                self.btn_export_hdf.setEnabled(True)

# For confirmation on closing the window
# def closeEvent(self, event):
#     reply = QMessageBox.question(self, 'Window Close', 'Are you sure you want to close the window?',
#                                  QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
#
#     if reply == QMessageBox.Yes:
#         event.accept()
#         print('Window closed')
#     else:
#         event.ignore()

# For converting into JSON and back
# df_apt_non_layer.to_json(
#     r'C:\Users\arjun\Downloads\APT_Code\APT_Project\df_apt_non_layer.json', orient='split')

# df_apt_non_layer = pd.read_json(r'C:\Users\arjun\Downloads\APT_Code\APT_Project\df_apt_non_layer.json',
#                                 orient='split')
